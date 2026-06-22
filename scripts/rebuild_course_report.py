from __future__ import annotations

import sys
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"E:\code\OA_SYSTEM")
sys.path.insert(0, str(ROOT / "scripts"))

import generate_project_documents as base  # noqa: E402


OUT = ROOT / "output" / "doc"
ASSETS = OUT / "report_rebuild_assets"
DIAGRAMS = ASSETS / "diagrams"
SCREENS = ASSETS / "screens"

FONT_HEI = r"C:\Windows\Fonts\simhei.ttf"
FONT_SONG = r"C:\Windows\Fonts\simsun.ttc"

REPORT_PATH = OUT / "SpringBoot微服务架构课程设计报告_OA_SYSTEM.docx"

PROJECT_CN = "企业 OA 管理系统"
PROJECT_EN = "OA_SYSTEM"
CLASS_NAME = "23gb软件1班"
TEACHER = "尹菠"
TEAM = [
    ("何成翔", "组长", "00000000000"),
    ("万宇轩", "组员", "00000000000"),
    ("刘灿", "组员", "00000000000"),
    ("姜周良", "组员", "00000000000"),
    ("吴发迟", "组员", "00000000000"),
]


TECH_STACK = [
    ("开发语言", "Java 17、TypeScript、HTML、SCSS"),
    ("后端框架", "Spring Boot 3.2.5、Spring Security、Spring WebSocket、Spring AMQP、Spring Mail"),
    ("持久层", "MyBatis-Plus 3.5.7、MySQL Connector/J、MySQL 8"),
    ("认证与权限", "JWT、BCrypt、Spring Security FilterChain、@PreAuthorize、Redis Token 黑名单"),
    ("缓存与消息", "Redis、RabbitMQ TopicExchange、手动 ACK 消费、站内信未读数缓存"),
    ("前端技术", "Vue 3、Vite 6、Element Plus、Pinia、Vue Router、Axios、ECharts"),
    ("文档与接口", "Knife4j / Swagger OpenAPI、统一 Result 返回结构"),
    ("部署依赖", "Docker Compose、MySQL、Redis、RabbitMQ、后端本地 upload 文件目录"),
]


TABLES = base.TABLES
MODULES = base.MODULES


COURSE_TOC = [
    (1, "绪论"),
    (1, "1 需求分析与设计"),
    (2, "1.1 系统需求分析"),
    (2, "1.2 系统设计方案"),
    (2, "1.3 功能分析"),
    (2, "1.4 数据流程分析"),
    (2, "1.5 业务流程分析"),
    (2, "1.6 系统实现技术选型"),
    (1, "2 系统设计"),
    (2, "2.1 系统概要设计"),
    (2, "2.2 系统结构设计"),
    (2, "2.3 数据库设计"),
    (1, "3 系统实现"),
    *[(2, f"3.{i} {m['title']}") for i, m in enumerate(MODULES, 1)],
    (1, "4 系统测试"),
    (2, "4.1 测试定义"),
    (2, "4.2 测试目的"),
    (2, "4.3 测试方法"),
    (2, "4.4 测试用例"),
    (2, "4.5 测试分析"),
    (1, "5 总结"),
]


SCREEN_SOURCE = {
    "login": ROOT / "frontend" / ".qa-full" / "01-login-desktop.png",
    "dashboard": ROOT / "frontend" / ".qa-full" / "02-dashboard-desktop.png",
    "users": ROOT / "frontend" / ".qa-full" / "users-departments-visible.png",
    "departments": ROOT / "frontend" / ".qa-screens" / "departments.png",
    "approval_list": ROOT / "frontend" / ".qa-screens" / "approvals.png",
    "approval_create": ROOT / "frontend" / ".qa-screens" / "approvals_create.png",
    "approval_todo": ROOT / "frontend" / ".qa-screens" / "approvals_todo.png",
    "schedules": ROOT / "frontend" / ".qa-screens" / "schedules.png",
    "calendar": ROOT / "frontend" / ".qa-screens" / "schedules_calendar.png",
    "news": ROOT / "frontend" / ".qa-screens" / "news-list-images.png",
    "news_manage": ROOT / "frontend" / ".qa-screens" / "news_manage.png",
    "dicts": ROOT / "frontend" / ".qa-screens" / "dicts.png",
    "profile": ROOT / "frontend" / ".qa-screens" / "profile-avatar-upload.png",
    "notifications": ROOT / "frontend" / ".qa-screens" / "notifications.png",
    "ai": ROOT / "frontend" / ".qa-full" / "06-ai-after-call.png",
}


SCREEN_CROPS = {
    "login": None,
    "dashboard": (180, 70, 1345, 1015),
    "users": (175, 75, 1340, 880),
    "departments": (175, 80, 1390, 850),
    "approval_list": (175, 80, 1390, 850),
    "approval_create": (175, 80, 1390, 850),
    "approval_todo": (175, 80, 1390, 850),
    "schedules": (175, 80, 1390, 850),
    "calendar": (175, 80, 1390, 850),
    "news": (175, 80, 1390, 850),
    "news_manage": (175, 80, 1390, 850),
    "dicts": (175, 80, 1390, 850),
    "profile": (175, 80, 1390, 850),
    "notifications": (175, 80, 1390, 850),
    "ai": (175, 80, 1345, 980),
}


MODULE_EXTRA = {
    "用户登录与认证权限模块": {
        "rule": "该模块以 AuthController 作为统一入口，以 UserServiceImpl 完成账号校验、注册、退出和 Token 刷新。登录成功后返回 AuthResponse，前端 Pinia 用户仓库保存 token 与用户资料，Axios 请求拦截器在后续请求中写入 Authorization 请求头。",
        "code": "关键实现点是 JwtAuthenticationFilter 在每次请求进入 Controller 前解析 Bearer Token，并将 LoginUser 写入 SecurityContext。SecurityConfig 中关闭 CSRF、启用无状态 Session，并为登录、注册、静态文件和 OpenAPI 文档放行。",
        "validation": "异常场景包括账号不存在、密码错误、账号禁用、Token 过期、Token 已加入 Redis 黑名单和权限不足。统一异常处理器把这些异常转换为规范 JSON 响应，前端根据 401 或 403 进行跳转或提示。",
    },
    "用户管理模块": {
        "rule": "用户管理只允许管理员操作，主要维护员工基本信息、所属部门、岗位、角色、启用状态和是否审批人。列表页按用户名、姓名、手机号、部门和状态组合查询，后端使用 MyBatis-Plus 分页对象返回 PageResult。",
        "code": "新增和重置密码时统一调用 BCryptPasswordEncoder 加密；删除使用逻辑删除字段，避免历史审批、日程和通知记录失去人员引用。更新审批人身份时只修改 is_approver 字段，不额外生成审批人表。",
        "validation": "系统没有独立角色表、岗位表和菜单表。角色以 sys_user.role 字段表示，岗位以 sys_user.position 字段保存，前端通过路由 meta.roles 控制菜单可见性，后端通过 Spring Security 注解控制接口权限。",
    },
    "部门组织管理模块": {
        "rule": "部门模块维护企业组织树，支持父子部门、负责人、默认审批人、排序和启停状态。审批提交时会读取申请人部门，优先匹配该部门 approver_id，缺省时再按业务规则处理。",
        "code": "DepartmentServiceImpl 负责树形结构组装和缓存刷新。部门树用于用户表单下拉、审批人匹配和部门管理页面展示，因此部门修改后需要清理 Redis 中的部门树缓存，避免页面仍显示旧组织关系。",
        "validation": "删除部门前需要注意是否已有用户归属或业务历史记录。当前 schema.sql 没有数据库外键约束，系统主要依靠 Service 层校验和逻辑删除维护数据一致性。",
    },
    "审批流程管理模块": {
        "rule": "审批模块覆盖请假、报销、加班和出差四类单据。申请人可创建草稿、修改草稿、提交、撤回；审批人可在待办列表中通过或驳回；管理员可以查看全量审批。",
        "code": "ApprovalServiceImpl 在 submit 方法中生成审批编号、设置 PENDING 状态、记录提交动作、解析审批人并发送通知。approve 和 reject 方法会校验当前用户是否具有审批权限，然后写入 oa_approval_record 并更新主表状态。",
        "validation": "源码中未集成 Activiti、Flowable 或 BPMN 引擎，因此本项目审批属于自研单级审批流。它的优点是结构轻量、业务路径清晰，适合课程设计；不足是多级会签、转办、加签等高级流程能力后续还需扩展。",
    },
    "日程会议管理模块": {
        "rule": "日程模块支持个人日程和会议日程。创建时可设置开始时间、结束时间、地点、提醒分钟数和参与人员；参与人可以接受或拒绝会议；首页和日历页按今天、本周和日期范围展示。",
        "code": "ScheduleReminderTask 定时扫描到期提醒，ScheduleServiceImpl 将需要提醒的日程转换为 NotificationMessage 投递到 RabbitMQ。消费者落库后通过 WebSocket 推送在线用户，并可触发邮件提醒。",
        "validation": "该模块同时体现了定时任务、消息队列、站内信和前端日历视图的协作关系。提醒状态由 oa_schedule.reminder_status 标识，避免同一日程重复推送。",
    },
    "新闻公告管理模块": {
        "rule": "新闻公告模块由管理员维护草稿、发布、下架和置顶，员工在新闻列表和详情页查看已发布内容，并可以评论、点赞、收藏。新闻详情读取后会增加阅读量，互动数据写入独立表。",
        "code": "NewsServiceImpl 对详情页启用 Redis 缓存，发布、下架、修改和删除时清理对应缓存。点赞和收藏使用 news_id 与 user_id 的唯一键避免重复记录，评论支持 parent_id 形成回复关系。",
        "validation": "AI 生成功能不是直接替代新闻管理，而是作为内容辅助能力存在。AI 返回的内容仍需管理员编辑、确认和发布，避免未经审核的内容直接进入公告区。",
    },
    "字典与基础数据管理模块": {
        "rule": "字典模块维护审批类型、审批状态、日程类型、新闻分类、通知类型和 AI 功能类型等基础选项。前端表单通过 typeCode 获取选项，减少硬编码，提高页面与后端枚举的一致性。",
        "code": "DictDataServiceImpl 优先读取 Redis 字典缓存，缓存不存在时查询 sys_dict_data 并写回 Redis。字典新增、修改、删除或手动刷新缓存后，前端再次读取即可获得最新选项。",
        "validation": "字典模块的真实表为 sys_dict_type 与 sys_dict_data。系统没有把所有枚举都完全数据库化，一些核心状态仍在 Java enum 中定义，文档中按源码实际情况区分说明。",
    },
    "文件上传与附件管理模块": {
        "rule": "文件模块统一处理头像、新闻封面、审批附件和日程附件等上传场景。上传接口接收 MultipartFile，并可附带 businessType 与 businessId，实现文件与业务数据的弱关联。",
        "code": "OaFileServiceImpl 按年月目录保存文件，使用 UUID 重命名避免重名覆盖，并把原始文件名、存储文件名、物理路径、访问 URL、文件大小、扩展名和上传人写入 sys_file。",
        "validation": "系统使用本地磁盘目录 backend/upload 和 /files 静态资源映射，没有接入 OSS、MinIO 或其他对象存储。报告中所有文件能力均按本地上传下载实现描述。",
    },
    "通知消息与实时推送模块": {
        "rule": "通知模块负责站内信列表、未读数量、标记已读、批量已读、删除和管理员系统通知。审批提交、审批处理、日程提醒和新闻发布都可以产生通知事件。",
        "code": "业务模块调用 NotificationProducer 发送消息到 RabbitMQ，NotificationConsumer 消费后调用 NotificationServiceImpl.createFromMessage 落库，再通过 NotificationWebSocketHandler 向在线用户推送。",
        "validation": "该模块把异步事件和实时界面连接起来。若 RabbitMQ 暂时不可用，通知投递链路会受影响；若用户不在线，站内信仍保存在 sys_notification，用户下次登录后可查看未读消息。",
    },
    "AI 助手与调用日志模块": {
        "rule": "AI 模块为审批摘要、审批风险分析、新闻生成、新闻润色、日程解析和智能问答提供统一入口。系统通过 oa.ai.provider 切换 mock 与 tongyi 实现，适合开发环境和真实服务环境分开使用。",
        "code": "AiController 只负责接收请求，具体能力由 AiService 接口实现。TongyiAiServiceImpl 通过 HTTP 调用通义千问接口，MockAiServiceImpl 返回可预测的模拟结果，每次调用都通过 AiLogServiceImpl 写入 sys_ai_log。",
        "validation": "AI 调用日志记录功能类型、模型提供商、模型名称、提示词、请求内容、响应内容、成功状态、错误信息和耗时。该表是源码中可确认的结构化日志表，不能扩写为完整操作审计日志系统。",
    },
}


MODULE_SCREEN = {
    1: [("login", "图 3.1 登录认证页面真实截图")],
    2: [("users", "图 3.2 用户管理页面真实截图")],
    3: [("departments", "图 3.3 部门管理页面真实截图")],
    4: [("approval_list", "图 3.4 审批列表页面真实截图"), ("approval_create", "图 3.5 审批创建页面真实截图"), ("approval_todo", "图 3.6 待办审批页面真实截图")],
    5: [("schedules", "图 3.7 日程列表页面真实截图"), ("calendar", "图 3.8 日程日历页面真实截图")],
    6: [("news", "图 3.9 新闻公告列表页面真实截图"), ("news_manage", "图 3.10 新闻公告管理页面真实截图")],
    7: [("dicts", "图 3.11 字典管理页面真实截图")],
    8: [("profile", "图 3.12 文件上传与个人资料页面真实截图")],
    9: [("notifications", "图 3.13 通知中心页面真实截图")],
    10: [("ai", "图 3.14 AI 助手调用页面真实截图")],
}


def ensure_dirs() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    SCREENS.mkdir(parents=True, exist_ok=True)


def ft(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_HEI if bold else FONT_SONG, size)


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    parts: list[str] = []
    for raw in text.split("\n"):
        line = ""
        for ch in raw:
            trial = line + ch
            if text_size(draw, trial, font)[0] <= max_width:
                line = trial
            else:
                if line:
                    parts.append(line)
                line = ch
        if line:
            parts.append(line)
    return parts or [""]


def center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.FreeTypeFont) -> None:
    lines = wrap_text(draw, text, font, box[2] - box[0] - 18)
    line_h = font.size + 8
    total_h = line_h * len(lines)
    y = box[1] + (box[3] - box[1] - total_h) // 2
    for line in lines:
        w, _ = text_size(draw, line, font)
        draw.text((box[0] + (box[2] - box[0] - w) // 2, y), line, fill="black", font=font)
        y += line_h


def rect(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, size: int = 26, bold: bool = True) -> None:
    draw.rectangle(box, fill="white", outline="black", width=3)
    center_text(draw, box, text, ft(size, bold))


def ellipse(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, size: int = 24) -> None:
    draw.ellipse(box, fill="white", outline="black", width=3)
    center_text(draw, box, text, ft(size, False))


def diamond(draw: ImageDraw.ImageDraw, cx: int, cy: int, w: int, h: int, text: str, size: int = 24) -> tuple[int, int, int, int]:
    pts = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
    draw.polygon(pts, fill="white", outline="black")
    draw.line([pts[0], pts[1], pts[2], pts[3], pts[0]], fill="black", width=3)
    center_text(draw, (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2), text, ft(size, True))
    return (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 3) -> None:
    import math

    draw.line([start, end], fill="black", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 13
    p1 = (end[0] - size * math.cos(angle - 0.45), end[1] - size * math.sin(angle - 0.45))
    p2 = (end[0] - size * math.cos(angle + 0.45), end[1] - size * math.sin(angle + 0.45))
    draw.polygon([end, p1, p2], fill="black")


def line(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 3) -> None:
    draw.line([start, end], fill="black", width=width)


def label(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, size: int = 22, bold: bool = False) -> None:
    draw.text(xy, text, fill="black", font=ft(size, bold))


def make_canvas(size: tuple[int, int]) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", size, "white")
    return img, ImageDraw.Draw(img)


def draw_er_diagram() -> Path:
    img, draw = make_canvas((1700, 1120))

    entities = {
        "dept": (70, 185, 320, 285, "sys_department\n部门"),
        "user": (470, 185, 720, 285, "sys_user\n用户"),
        "approval": (880, 185, 1130, 285, "oa_approval\n审批"),
        "record": (1320, 185, 1620, 285, "oa_approval_record\n审批记录"),
        "notification": (70, 555, 340, 655, "sys_notification\n通知"),
        "schedule": (470, 555, 720, 655, "oa_schedule\n日程"),
        "participant": (880, 555, 1210, 655, "oa_schedule_participant\n日程参与人"),
        "dict_type": (70, 885, 340, 985, "sys_dict_type\n字典类型"),
        "dict_data": (470, 885, 740, 985, "sys_dict_data\n字典数据"),
        "news": (880, 885, 1130, 985, "oa_news\n新闻公告"),
        "file": (1320, 885, 1620, 985, "sys_file\n文件"),
    }

    for box in entities.values():
        rect(draw, box[:4], box[4], 25, True)

    def attr(entity_key: str, box: tuple[int, int, int, int], text: str, anchor: str = "top") -> None:
        ellipse(draw, box, text, 20)
        ent = entities[entity_key]
        if anchor == "top":
            target = ((ent[0] + ent[2]) // 2, ent[1])
            source = ((box[0] + box[2]) // 2, box[3])
        elif anchor == "bottom":
            target = ((ent[0] + ent[2]) // 2, ent[3])
            source = ((box[0] + box[2]) // 2, box[1])
        elif anchor == "left":
            target = (ent[0], (ent[1] + ent[3]) // 2)
            source = (box[2], (box[1] + box[3]) // 2)
        else:
            target = (ent[2], (ent[1] + ent[3]) // 2)
            source = (box[0], (box[1] + box[3]) // 2)
        line(draw, source, target, 2)

    attr("dept", (55, 60, 205, 120), "id(PK)")
    attr("dept", (215, 60, 380, 120), "name")
    attr("user", (405, 60, 555, 120), "id(PK)")
    attr("user", (565, 60, 765, 120), "username")
    attr("approval", (830, 60, 1080, 120), "approval_no(PK)")
    attr("approval", (1088, 60, 1245, 120), "status")
    attr("schedule", (455, 430, 610, 490), "title")
    attr("schedule", (620, 430, 840, 490), "start_time")
    attr("news", (835, 770, 990, 830), "title")
    attr("news", (1005, 770, 1170, 830), "status")
    attr("notification", (55, 690, 220, 750), "title", "bottom")
    attr("notification", (230, 690, 435, 750), "read_status", "bottom")

    def h_relation(left_key: str, right_key: str, name: str, card_l: str, card_r: str) -> None:
        left = entities[left_key]
        right = entities[right_key]
        cy = (left[1] + left[3]) // 2
        cx = (left[2] + right[0]) // 2
        diamond(draw, cx, cy, 150, 90, name, 22)
        line(draw, (left[2], cy), (cx - 75, cy), 3)
        line(draw, (cx + 75, cy), (right[0], cy), 3)
        label(draw, (left[2] + 20, cy - 42), card_l, 23, True)
        label(draw, (right[0] - 45, cy - 42), card_r, 23, True)

    def v_relation(top_key: str, bottom_key: str, name: str, card_t: str, card_b: str) -> None:
        top = entities[top_key]
        bottom = entities[bottom_key]
        cx = (top[0] + top[2]) // 2
        cy = (top[3] + bottom[1]) // 2
        diamond(draw, cx, cy, 150, 90, name, 22)
        line(draw, (cx, top[3]), (cx, cy - 45), 3)
        line(draw, (cx, cy + 45), (cx, bottom[1]), 3)
        label(draw, (cx + 35, top[3] + 35), card_t, 23, True)
        label(draw, (cx + 35, bottom[1] - 65), card_b, 23, True)

    h_relation("dept", "user", "拥有", "1", "N")
    h_relation("user", "approval", "提交", "1", "N")
    h_relation("approval", "record", "产生", "1", "N")
    v_relation("user", "schedule", "创建", "1", "N")
    h_relation("schedule", "participant", "邀请", "1", "N")
    h_relation("dict_type", "dict_data", "包含", "1", "N")
    h_relation("news", "file", "关联", "1", "N")

    diamond(draw, 205, 420, 150, 90, "接收", 22)
    line(draw, (470, 245), (280, 420), 3)
    line(draw, (130, 420), (205, 555), 3)
    label(draw, (320, 300), "1", 23, True)
    label(draw, (145, 505), "N", 23, True)

    path = DIAGRAMS / "er_blackwhite.png"
    img.save(path)
    return path


def draw_operation_flow() -> Path:
    img, draw = make_canvas((1000, 1260))
    rect(draw, (390, 45, 610, 115), "登录系统", 25)
    arrow(draw, (500, 115), (500, 165))
    rect(draw, (350, 165, 650, 235), "进入工作台", 25)
    arrow(draw, (500, 235), (500, 295))
    rect(draw, (330, 295, 670, 365), "选择业务模块", 25)
    arrow(draw, (500, 365), (500, 430))
    diamond(draw, 500, 500, 210, 150, "操作类型", 25)

    operations = [
        ("新增", (60, 660, 240, 730), "填写业务表单"),
        ("修改", (290, 660, 470, 730), "选择记录并编辑"),
        ("删除", (530, 660, 710, 730), "确认删除或留痕"),
        ("查询", (760, 660, 940, 730), "输入条件筛选"),
    ]
    for name, box, text in operations:
        arrow(draw, (500, 500), ((box[0] + box[2]) // 2, box[1]))
        label(draw, ((box[0] + box[2]) // 2 - 25, box[1] - 45), name, 23)
        rect(draw, box, text, 22)
        arrow(draw, ((box[0] + box[2]) // 2, box[3]), ((box[0] + box[2]) // 2, 830))

    rect(draw, (350, 830, 650, 900), "后端参数校验", 24)
    arrow(draw, (500, 900), (500, 960))
    diamond(draw, 500, 1030, 230, 140, "校验通过?", 24)
    arrow(draw, (500, 1100), (500, 1160))
    rect(draw, (350, 1160, 650, 1230), "写入数据库并返回结果", 22)
    arrow(draw, (615, 1030), (850, 1030))
    label(draw, (720, 990), "否", 23)
    rect(draw, (770, 995, 940, 1065), "返回错误提示", 22)
    arrow(draw, (385, 1030), (160, 1030))
    label(draw, (235, 990), "是", 23)
    rect(draw, (60, 995, 230, 1065), "刷新页面数据", 22)

    path = DIAGRAMS / "operation_flow_blackwhite.png"
    img.save(path)
    return path


def draw_use_case() -> Path:
    img, draw = make_canvas((1300, 940))

    def actor(cx: int, cy: int, name: str) -> None:
        draw.ellipse((cx - 18, cy - 60, cx + 18, cy - 24), outline="black", width=3)
        line(draw, (cx, cy - 24), (cx, cy + 55), 3)
        line(draw, (cx - 45, cy + 5), (cx + 45, cy + 5), 3)
        line(draw, (cx, cy + 55), (cx - 45, cy + 115), 3)
        line(draw, (cx, cy + 55), (cx + 45, cy + 115), 3)
        label(draw, (cx - 55, cy + 130), name, 31, True)

    actor(130, 135, "管理员")
    actor(130, 425, "员工")
    actor(130, 700, "审批人员")

    use_cases = [
        ("用户管理", 450, 70, 720, 145, [130]),
        ("部门管理", 760, 70, 1030, 145, [130]),
        ("字典管理", 1035, 180, 1260, 255, [130]),
        ("新闻公告管理", 455, 205, 760, 280, [130]),
        ("系统通知管理", 780, 320, 1085, 395, [130]),
        ("登录认证", 430, 420, 670, 495, [420]),
        ("个人信息维护", 710, 430, 1030, 505, [420]),
        ("发起审批", 430, 550, 690, 625, [420]),
        ("日程会议管理", 730, 560, 1050, 635, [420]),
        ("查看新闻与互动", 430, 700, 780, 775, [420]),
        ("站内信处理", 820, 700, 1120, 775, [420, 700]),
        ("待办审批处理", 430, 825, 760, 900, [700]),
        ("审批记录查看", 800, 825, 1120, 900, [700]),
        ("AI 辅助办公", 1030, 520, 1260, 595, [420, 700]),
    ]

    for text, x1, y1, x2, y2, actors in use_cases:
        ellipse(draw, (x1, y1, x2, y2), text, 27)
        for ay in actors:
            start = (175, ay)
            end = (x1, (y1 + y2) // 2)
            arrow(draw, start, end, 2)

    path = DIAGRAMS / "use_case_blackwhite.png"
    img.save(path)
    return path


def draw_function_structure() -> Path:
    img, draw = make_canvas((1700, 980))
    top = (625, 50, 1075, 125)
    rect(draw, top, "企业 OA 管理系统", 29)
    groups = [
        ("组织权限", 145, 225, ["登录认证", "用户管理", "部门管理", "权限控制"]),
        ("办公业务", 545, 225, ["审批流程", "日程会议", "新闻公告", "文件附件"]),
        ("基础数据", 945, 225, ["字典类型", "字典数据", "统一返回", "异常处理"]),
        ("消息智能", 1345, 225, ["站内信通知", "WebSocket推送", "RabbitMQ队列", "AI助手"]),
    ]
    for group, gx, gy, leaves in groups:
        rect(draw, (gx, gy, gx + 210, gy + 70), group, 25)
        line(draw, ((top[0] + top[2]) // 2, top[3]), (gx + 105, gy), 3)
        start_x = gx - 80
        for i, leaf in enumerate(leaves):
            x = start_x + i * 95
            y = gy + 170
            line(draw, (gx + 105, gy + 70), (x + 35, y), 2)
            draw.rectangle((x, y, x + 70, y + 420), fill="white", outline="black", width=3)
            chars = "\n".join(list(leaf))
            center_text(draw, (x, y, x + 70, y + 420), chars, ft(25, True))

    path = DIAGRAMS / "function_structure_blackwhite.png"
    img.save(path)
    return path


def make_diagrams() -> dict[str, Path]:
    return {
        "er": draw_er_diagram(),
        "operation": draw_operation_flow(),
        "use_case": draw_use_case(),
        "function": draw_function_structure(),
    }


def preprocess_screens() -> dict[str, Path]:
    result: dict[str, Path] = {}
    for key, src in SCREEN_SOURCE.items():
        if not src.exists():
            continue
        img = Image.open(src).convert("RGB")
        crop = SCREEN_CROPS.get(key)
        if crop:
            w, h = img.size
            x1, y1, x2, y2 = crop
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(w, x2), min(h, y2)
            img = img.crop((x1, y1, x2, y2))
        img = ImageOps.expand(img, border=2, fill="#333333")
        if img.width < 1700:
            scale = 1700 / img.width
            img = img.resize((int(img.width * scale), int(img.height * scale)), Image.Resampling.LANCZOS)
        out = SCREENS / f"{key}.png"
        img.save(out)
        result[key] = out
    return result


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, name: str = "宋体") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(3)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        st = doc.styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.paragraph_format.first_line_indent = Pt(0)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(6)


def add_p(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r)


def add_no_indent(doc: Document, text: str, size: float = 10.5, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(h)
        set_run_font(r, 10, True, name="黑体")
        shade_cell(hdr[i], "FFFFFF")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, 9.5)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(item)
        set_run_font(r)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(r, 10)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.15) -> None:
    if not path.exists():
        add_p(doc, f"未找到图片文件：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    add_no_indent(doc, "SpringBoot 微服务架构课程设计报告", 22, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_no_indent(doc, PROJECT_CN, 20, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_no_indent(doc, PROJECT_EN, 14, False, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        doc.add_paragraph()
    add_table(
        doc,
        ["项目项", "内容", "项目项", "内容"],
        [
            ("项目名称", PROJECT_CN, "系统标识", PROJECT_EN),
            ("班级", CLASS_NAME, "指导老师", TEACHER),
            ("组长", "何成翔", "日期", ""),
        ],
        [3, 5, 3, 5],
    )
    add_table(doc, ["姓名", "小组职责", "学号"], TEAM, [4, 5, 5])
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_no_indent(doc, "目  录", 16, True, WD_ALIGN_PARAGRAPH.CENTER)
    for level, text in COURSE_TOC:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt((level - 1) * 20)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        set_run_font(r, 10.5, level == 1)
    doc.add_page_break()


def add_intro(doc: Document, diagrams: dict[str, Path], screens: dict[str, Path]) -> None:
    doc.add_heading("绪论", 1)
    add_p(doc, "企业日常办公活动涉及人员资料维护、部门协同、审批流转、会议安排、公告发布和消息通知等多类信息。如果仍依赖线下表单、即时通讯零散传递和人工统计，不仅会增加管理成本，也容易造成审批记录缺失、通知触达不及时、文件附件难以归档等问题。")
    add_p(doc, "OA_SYSTEM 面向企业内部办公管理场景，使用 Spring Boot 与 Vue3 构建前后端分离系统，将用户、部门、审批、日程、新闻公告、字典、文件、通知和 AI 助手整合到统一平台中。系统的重点不是制作展示型页面，而是围绕真实办公流程完成数据录入、状态流转、权限控制和消息触达。")
    add_p(doc, "经本地源码核查，本项目后端是一个单体 Spring Boot 模块化应用，并不是由多个独立后端服务组成的 Spring Cloud 集群。因此，本报告在课程主题“SpringBoot 微服务架构”下重点说明模块化分层、缓存、消息队列、实时通信和容器化依赖协作，不把源码中不存在的网关、注册中心、配置中心、Feign 远程调用或 Sentinel 限流写成已实现内容。")
    add_p(doc, "本报告按照课程设计报告模板展开，第一部分进行需求分析与设计，第二部分说明系统设计和数据库设计，第三部分按照真实源码模块逐一说明系统实现，第四部分给出测试方法与测试用例，第五部分总结项目完成情况和后续改进方向。")
    add_image(doc, diagrams["function"], "图 0.1 系统功能结构图")
    if "dashboard" in screens:
        add_image(doc, screens["dashboard"], "图 0.2 系统首页工作台真实截图", 6.25)


def add_requirement_chapter(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("1 需求分析与设计", 1)
    doc.add_heading("1.1 系统需求分析", 2)
    add_p(doc, "从业务角度看，系统需要解决三个核心问题：第一，企业员工和部门信息需要集中维护，避免审批人、部门负责人、岗位信息分散；第二，审批、会议、公告和通知需要在线流转并保留记录；第三，管理员需要通过统一后台完成基础数据维护，员工需要通过统一入口处理个人办公事项。")
    add_p(doc, "从角色角度看，系统主要面向管理员、普通员工和审批人员。管理员负责组织权限和基础数据维护，普通员工负责个人办公业务，审批人员负责待办审批处理。审批人员不是独立角色表中的角色，而是 sys_user.is_approver 字段标识出的员工身份。")
    add_table(
        doc,
        ["用户角色", "主要业务需求", "权限边界"],
        [
            ("管理员", "维护用户、部门、字典、新闻公告、系统通知，查看审批和 AI 调用日志。", "前端菜单通过 roles 控制，后端接口通过 hasRole('ADMIN') 控制。"),
            ("普通员工", "登录系统、维护个人资料、发起审批、查看公告、管理日程、接收站内信、使用 AI 助手。", "只能访问本人创建、本人接收或本人参与的数据。"),
            ("审批人员", "在员工能力基础上处理待本人审批的审批单，查看审批记录并使用审批 AI 辅助能力。", "通过 is_approver 与审批人字段共同判断，不存在单独审批角色表。"),
        ],
        [3, 7, 6],
    )
    add_p(doc, "从非功能角度看，系统应保证接口返回结构统一、异常提示清晰、密码加密保存、Token 具备过期与退出失效能力、常用字典和未读数量能够缓存、通知类业务能够异步处理、页面在常用浏览器中正常显示。")

    doc.add_heading("1.2 系统设计方案", 2)
    add_p(doc, "系统采用前后端分离设计。前端 Vue3 应用负责路由、页面、表单、状态和交互反馈；后端 Spring Boot 应用负责认证授权、业务规则、事务处理、数据库访问、缓存读写、消息投递、文件保存和 AI 服务调用。前后端通过 RESTful JSON API 交互，通知消息通过 WebSocket 补充实时通道。")
    add_p(doc, "在数据存储方面，MySQL 作为主数据库保存业务实体；Redis 用于字典、部门树、新闻详情和未读通知数量等高频读取数据；RabbitMQ 用于审批、日程、新闻和系统通知的异步事件；本地 upload 目录保存头像、封面图和附件等文件。")
    add_p(doc, "在部署方面，项目提供 docker-compose.yml 启动 MySQL、Redis 和 RabbitMQ 等依赖服务，后端应用默认端口为 8081，前端开发环境由 Vite 提供页面服务。实际运行时需要注意 application.yml 中 RabbitMQ 端口与 docker-compose.yml 暴露端口是否一致。")
    add_table(doc, ["设计层次", "设计内容"], [
        ("表现层", "Vue3 单页应用、Element Plus 组件、Pinia 状态、Vue Router 路由、Axios 请求拦截。"),
        ("接口层", "Controller 接收请求并返回 Result，参数通过 DTO 校验，OpenAPI/Knife4j 生成接口文档。"),
        ("业务层", "Service 处理事务、权限、状态流转、缓存、消息和文件等业务逻辑。"),
        ("数据层", "Mapper 继承 BaseMapper，实体类通过 @TableName 映射数据库表。"),
        ("基础设施层", "MySQL、Redis、RabbitMQ、WebSocket、Spring Mail、本地文件目录。"),
    ], [4, 12])

    doc.add_heading("1.3 功能分析", 2)
    add_p(doc, "根据源码目录、前端路由和数据库表结构，OA_SYSTEM 的功能可以划分为组织权限、办公业务、基础数据、消息智能四个功能域。每个功能域下面由多个真实模块组成，模块之间通过用户、部门、审批、通知等数据关系协同工作。")
    add_image(doc, diagrams["use_case"], "图 1.1 系统用例图")
    add_table(doc, ["功能模块", "功能范围"], [(m["title"], m["functions"]) for m in MODULES], [4.6, 11.4])

    doc.add_heading("1.4 数据流程分析", 2)
    add_p(doc, "系统数据流从用户登录开始。登录成功后，前端保存 JWT，并在每次请求中附带 Authorization 请求头；后端过滤器解析 Token 后写入安全上下文，Controller 才能取得当前用户身份并进入业务处理。")
    add_p(doc, "普通业务数据流遵循“页面表单或查询条件 - API 请求 - Controller - Service - Mapper - MySQL - Result - 页面渲染”的路径。对于字典、部门树、新闻详情和未读数等高频读取数据，Service 会优先查询 Redis，缓存未命中后再访问数据库。")
    add_p(doc, "通知类数据流存在异步处理。审批提交、审批处理、日程提醒、系统通知等业务先构造 NotificationMessage 并投递到 RabbitMQ，消费者落库 sys_notification 后再尝试 WebSocket 在线推送，未在线用户仍可在通知中心查看历史消息。")
    add_image(doc, diagrams["er"], "图 1.2 系统 E-R 图")

    doc.add_heading("1.5 业务流程分析", 2)
    add_p(doc, "系统常见操作流程具有一致性：用户登录后进入工作台，选择具体业务模块，根据操作类型执行新增、修改、删除、查询或审批处理；后端完成参数校验、权限校验和业务校验后写入数据库并返回结果。")
    add_p(doc, "审批流程是系统中最具代表性的业务流程。员工先创建草稿，确认信息后提交；系统生成审批编号并匹配部门审批人；审批人处理后，系统更新审批状态、写入审批记录并向申请人发送通知。")
    add_image(doc, diagrams["operation"], "图 1.3 系统操作流程图", 4.2)

    doc.add_heading("1.6 系统实现技术选型", 2)
    add_p(doc, "本项目技术选型以课程要求和本地源码为准。后端采用 Spring Boot 3.2.5 和 Java 17，前端采用 Vue3 与 Vite，数据库使用 MySQL 8，缓存使用 Redis，异步消息使用 RabbitMQ，页面组件使用 Element Plus。")
    add_table(doc, ["类别", "实际技术选型"], TECH_STACK, [4.5, 11.5])
    add_p(doc, "Spring Boot 的自动配置和注解式开发降低了工程搭建复杂度，Controller-Service-Mapper 分层使业务代码具有清晰边界。Spring Security 与 JWT 适合前后端分离场景，可以避免传统 Session 对后端横向扩展的限制。")
    add_p(doc, "MyBatis-Plus 提供通用 CRUD、分页和 LambdaQueryWrapper，减少重复 SQL 编写。Redis 适用于保存短期缓存和 Token 黑名单，RabbitMQ 适合把通知生成和业务提交解耦，WebSocket 则负责把站内信事件实时推送给在线用户。")


def add_design_chapter(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("2 系统设计", 1)
    doc.add_heading("2.1 系统概要设计", 2)
    add_p(doc, "系统整体采用模块化单体架构。后端以 com.example.oa.module 作为业务模块根目录，每个模块内部再划分 controller、service、mapper、entity、dto、vo、enums 等包；公共能力放在 common、config、security 中，避免业务模块重复实现基础逻辑。")
    add_p(doc, "前端采用管理端单页应用结构，src/views 存放业务页面，src/api 存放请求封装，src/stores 存放用户、字典和通知等状态，src/router 统一维护页面路由和权限元信息。")
    add_table(doc, ["层次", "后端职责", "前端对应"], [
        ("表现层", "无页面模板，仅提供 JSON API。", "Vue 页面、Element Plus 组件、表格、表单、弹窗和图表。"),
        ("接口层", "Controller 定义路径、方法和参数校验。", "api/*.ts 封装请求路径和参数。"),
        ("业务层", "ServiceImpl 处理事务、状态流转、缓存、消息和权限。", "页面组合式逻辑调用接口并处理返回。"),
        ("数据层", "Mapper 与 Entity 映射数据库表。", "VO/DTO 数据在页面中展示或编辑。"),
        ("公共层", "Result、PageResult、GlobalExceptionHandler、JwtAuthenticationFilter。", "Axios 拦截器、路由守卫、Pinia 状态。"),
    ], [3, 7, 6])

    doc.add_heading("2.2 系统结构设计", 2)
    add_p(doc, "后端启动入口为 OaSystemApplication，配置文件 application.yml 定义服务端口、数据库、Redis、RabbitMQ、文件上传路径、AI 提供商和 OpenAPI 等参数。项目并未配置 Spring Cloud Gateway、Nacos、Eureka 或 Config Server，因此结构设计以单应用内部模块协作为主。")
    add_table(doc, ["源码目录", "主要内容"], [
        ("backend/src/main/java/com/example/oa/common", "统一返回 Result、分页 PageResult、异常处理、枚举、基础实体、安全上下文工具。"),
        ("backend/src/main/java/com/example/oa/config", "MyBatis-Plus、Redis、RabbitMQ、WebSocket、静态文件访问、OpenAPI、异步任务配置。"),
        ("backend/src/main/java/com/example/oa/security", "Spring Security 配置、JWT 工具、认证过滤器、登录用户和权限判断。"),
        ("backend/src/main/java/com/example/oa/module/user", "登录注册、用户管理、个人资料、审批人标识。"),
        ("backend/src/main/java/com/example/oa/module/approval", "审批单、审批记录、审批状态流转、AI 摘要和风险分析。"),
        ("backend/src/main/java/com/example/oa/module/notification", "站内信、RabbitMQ 生产消费、WebSocket 推送、邮件提醒。"),
        ("frontend/src/views", "登录、工作台、用户、部门、审批、日程、新闻、字典、通知、AI 和个人中心页面。"),
    ], [6.2, 9.8])
    add_image(doc, diagrams["function"], "图 2.1 系统功能结构图")

    doc.add_heading("2.3 数据库设计", 2)
    add_p(doc, "数据库脚本位于 backend/src/main/resources/db/schema.sql，数据库名为 oa_management。脚本共定义 15 张业务表，覆盖用户、部门、审批、日程、新闻、字典、文件、通知和 AI 调用日志。")
    add_p(doc, "schema.sql 中没有显式定义数据库外键约束，业务关联主要依赖字段命名和 Service 层逻辑维护。例如 oa_approval.applicant_id 指向 sys_user.id，oa_approval.department_id 指向 sys_department.id，oa_schedule_participant.schedule_id 指向 oa_schedule.id。")
    add_image(doc, diagrams["er"], "图 2.2 系统 E-R 图")
    add_table(doc, ["数据表", "中文说明", "关键字段或用途"], TABLES, [4, 4, 8])
    add_p(doc, "数据库设计中大量表包含 create_time、update_time、create_by、update_by、deleted 等公共字段，体现了通用审计字段和逻辑删除设计。对于点赞、收藏和日程参与人等关系表，脚本设置唯一键防止重复关系。")


def add_module_section(doc: Document, idx: int, module: dict[str, str], screens: dict[str, Path]) -> None:
    title = module["title"]
    extra = MODULE_EXTRA.get(title, {})
    doc.add_heading(f"3.{idx} {title}", 2)

    doc.add_heading(f"3.{idx}.1 模块功能说明", 3)
    add_p(doc, module["functions"])
    add_p(doc, extra.get("rule", "该模块按照 Controller、Service、Mapper 和 Entity 分层实现，接口返回统一 Result，页面通过 Axios 调用对应后端 API。"))
    add_p(doc, "从课程设计实现角度看，本模块不仅需要完成页面上的增删改查，还要保证数据权限、状态校验、异常返回和前后端字段一致。报告中所列功能均依据本地源码、前端路由和数据库脚本整理。")

    doc.add_heading(f"3.{idx}.2 业务流程说明", 3)
    add_p(doc, module["flow"])
    add_p(doc, "流程中的每一次状态变化都由 Service 层统一处理，Controller 不直接拼接业务规则。这样可以避免相同操作在不同接口中出现逻辑分叉，也便于后续补充事务、缓存清理或消息通知。")
    add_p(doc, "前端页面通常先完成基础表单校验，再向后端提交请求；后端仍会进行必要校验，防止绕过前端页面直接调用接口造成无效数据写入。")

    doc.add_heading(f"3.{idx}.3 核心类、接口与方法", 3)
    add_p(doc, module["classes"])
    add_p(doc, extra.get("code", "核心实现采用 MyBatis-Plus 提供的 BaseMapper 和 ServiceImpl，并结合 LambdaQueryWrapper 进行条件查询。涉及状态变更的接口由 Service 层封装，避免 Controller 直接修改实体状态。"))
    add_table(doc, ["实现项", "具体说明"], [
        ("核心类/接口", module["classes"]),
        ("前端页面或组件", module["frontend"]),
        ("主要接口", module["api"]),
    ], [3.8, 12.2])

    doc.add_heading(f"3.{idx}.4 数据库表与关键字段", 3)
    add_p(doc, module["tables"])
    add_p(doc, "数据库字段设计与实体类通过 @TableName 和字段名进行映射。对于分页查询，前端传入 pageNum、pageSize 和筛选条件，后端返回记录列表、总数、当前页和页大小。")
    add_p(doc, extra.get("validation", "本模块的数据一致性主要由 Service 层校验保证，数据库脚本未对所有业务关系设置外键约束。"))

    doc.add_heading(f"3.{idx}.5 前后端交互逻辑", 3)
    add_p(doc, module["frontend"])
    add_p(doc, module["api"])
    add_p(doc, "前端通过 src/api 下的封装函数调用后端接口，页面组件只关注表单、表格和操作反馈。后端返回 Result.success 或业务异常，前端请求拦截器统一处理登录失效和权限不足。")

    doc.add_heading(f"3.{idx}.6 页面效果与实现小结", 3)
    for key, caption in MODULE_SCREEN.get(idx, []):
        if key in screens:
            add_image(doc, screens[key], caption, 6.25)
    add_p(doc, module["summary"])
    add_p(doc, "从实现完整性看，该模块已经具备课程设计验收所需的核心功能。后续若要继续工程化，可以补充更细粒度权限、操作审计日志、自动化接口测试和更多异常场景测试。")


def add_implementation_chapter(doc: Document, screens: dict[str, Path]) -> None:
    doc.add_heading("3 系统实现", 1)
    add_p(doc, "本章按照项目源码中的真实业务模块展开。每个模块均从功能说明、业务流程、核心类和接口、数据库字段、前后端交互、关键实现思路、页面效果和模块小结等方面说明，避免只停留在概述层面。")
    add_p(doc, "课程要求中提到的用户登录、权限管理、用户管理、角色/菜单/部门/岗位、通知公告、流程审批、系统日志、文件上传、网关路由、服务注册发现和微服务调用等内容，需要按项目真实情况处理。本项目实现了登录认证、轻量角色权限、用户、部门、审批、日程、新闻公告、字典、文件、通知和 AI 日志；未实现独立角色表、菜单表、岗位表、网关、注册中心、配置中心和跨微服务调用。")
    for idx, module in enumerate(MODULES, 1):
        add_module_section(doc, idx, module, screens)


def add_test_chapter(doc: Document) -> None:
    doc.add_heading("4 系统测试", 1)
    doc.add_heading("4.1 测试定义", 2)
    add_p(doc, "系统测试是对系统功能、接口、权限、数据一致性、构建结果和页面交互进行综合验证的过程。对于本项目而言，测试不仅包括页面能否正常显示，还包括后端接口是否能正确识别用户身份、业务状态是否按规则流转、缓存和消息队列是否与业务结果一致。")
    add_p(doc, "由于项目为前后端分离系统，测试需要同时覆盖后端 Maven 构建、前端 Vite 构建、数据库初始化脚本、API 路径、前端路由和真实页面截图。")

    doc.add_heading("4.2 测试目的", 2)
    add_p(doc, "测试目的包括：验证用户能够正常登录并携带 Token 访问接口；验证管理员与员工权限边界；验证用户、部门、审批、日程、新闻、字典、文件、通知和 AI 模块的主要业务路径；验证异常处理和统一返回结构；验证项目能够完成构建。")
    add_p(doc, "同时，测试还用于发现配置层面的风险。例如 application.yml 中 RabbitMQ 端口和 docker-compose.yml 暴露端口存在差异，若运行环境没有同步修改，将影响通知队列链路。")

    doc.add_heading("4.3 测试方法", 2)
    add_bullets(doc, [
        "源码结构测试：检查 Controller、Service、Mapper、Entity、DTO、VO 和前端路由是否与报告模块一致。",
        "数据库脚本测试：检查 schema.sql 与 data.sql 是否能覆盖报告中的表结构、字段和初始化数据。",
        "构建测试：后端执行 mvn -q -DskipTests package，前端执行 npm run build。",
        "权限测试：使用管理员和员工身份分别访问管理接口与普通业务接口。",
        "流程测试：围绕审批提交、审批处理、日程提醒、新闻发布和站内信推送设计端到端用例。",
        "页面测试：使用真实项目截图检查页面结构、菜单、表格、表单和操作按钮是否与源码功能一致。",
    ])

    doc.add_heading("4.4 测试用例", 2)
    add_table(doc, ["编号", "测试模块", "测试步骤", "预期结果"], [
        ("TC-01", "登录认证", "输入 data.sql 中的 admin 或 employee 账号密码，调用 /api/auth/login。", "返回 token 和用户信息，密码不明文返回。"),
        ("TC-02", "Token 鉴权", "登录后调用 /api/auth/me，退出后再次访问受保护接口。", "登录状态返回用户信息，退出后 Token 加入黑名单并被拒绝。"),
        ("TC-03", "权限控制", "员工访问 /api/users/page、/api/dict-types/page 等管理员接口。", "接口返回权限不足或页面不展示对应菜单。"),
        ("TC-04", "用户管理", "管理员新增用户、修改资料、禁用账号、重置密码。", "sys_user 对应字段正确变化，密码字段为 BCrypt 密文。"),
        ("TC-05", "部门管理", "新增子部门、设置负责人和审批人，再查看部门树。", "部门树层级正确，部门缓存刷新后前端选项更新。"),
        ("TC-06", "审批流程", "员工创建请假审批并提交，审批人通过或驳回。", "oa_approval 状态更新，oa_approval_record 新增记录，通知发送。"),
        ("TC-07", "日程会议", "创建会议日程并添加参与人，参与人接受或拒绝。", "oa_schedule 和 oa_schedule_participant 数据正确，日历页可查询。"),
        ("TC-08", "新闻公告", "管理员发布新闻，员工查看、评论、点赞和收藏。", "新闻状态为 PUBLISHED，互动表和计数字段更新。"),
        ("TC-09", "字典缓存", "修改字典项后刷新缓存，再在表单中读取选项。", "前端选项与 sys_dict_data 最新数据一致。"),
        ("TC-10", "文件上传", "上传头像或新闻封面，访问返回 file_url。", "sys_file 写入元数据，本地 upload 目录存在文件。"),
        ("TC-11", "通知推送", "提交审批或发送系统通知，观察通知中心。", "sys_notification 新增记录，未读数更新，在线用户可收到 WebSocket 消息。"),
        ("TC-12", "AI 助手", "调用智能问答、审批摘要或日程解析接口。", "返回 AiResponse，并在 sys_ai_log 中记录调用日志。"),
    ], [2, 3, 7, 6])

    doc.add_heading("4.5 测试分析", 2)
    add_p(doc, "根据本次本地检查，后端 Maven 打包和前端 Vite 构建均可作为基础质量验证手段。源码结构、数据库脚本和前端路由能够支撑报告中描述的真实功能模块。")
    add_p(doc, "测试过程中需要重点关注三类风险：第一，RabbitMQ 运行端口配置是否与本地容器一致；第二，通义千问 API Key 是否已按真实运行环境配置；第三，页面截图只反映已有 QA 环境和初始化数据，正式验收时应以现场启动后的真实数据为准。")
    add_p(doc, "从课程设计角度看，系统已经具备登录认证、权限控制、基础资料维护、核心办公业务、通知推送和 AI 辅助能力。尚未实现的 Spring Cloud 网关、注册中心、配置中心、跨服务调用和完整操作审计日志，在报告中已按未实现内容说明，没有作为已完成功能描述。")


def add_summary(doc: Document) -> None:
    doc.add_heading("5 总结", 1)
    add_p(doc, "本课程设计围绕企业 OA 办公场景，完成了 OA_SYSTEM 的需求分析、系统设计、数据库设计、模块实现和测试说明。系统以 Spring Boot 3.2.5 为后端基础，以 Vue3 为前端基础，结合 MySQL、Redis、RabbitMQ、WebSocket 和 AI 服务接入，实现了较完整的企业办公管理流程。")
    add_p(doc, "在功能层面，系统覆盖了用户登录与认证、用户管理、部门管理、审批流程、日程会议、新闻公告、字典管理、文件上传、通知推送和 AI 助手。各模块之间不是孤立存在的，例如审批提交依赖用户和部门数据，审批状态变化会产生通知消息，日程提醒会经过 RabbitMQ 和 WebSocket 推送，新闻封面和头像依赖文件上传模块。")
    add_p(doc, "在技术层面，系统采用前后端分离、统一返回、统一异常处理、JWT 鉴权、逻辑删除、分页查询、Redis 缓存、RabbitMQ 异步消息和 WebSocket 实时推送等实现方式。虽然项目名称对应课程中的微服务架构主题，但当前源码没有拆分为多个独立微服务，报告已经严格按实际项目情况说明。")
    add_p(doc, "项目仍有可提升空间。后续可以在现有模块化单体基础上进一步补充操作审计日志、角色菜单权限表、岗位表、接口自动化测试、部署脚本、RabbitMQ 死信队列、文件对象存储和更完善的多级审批流程。如果课程后续要求真正的 Spring Cloud 微服务，也可以优先拆分认证服务、通知服务、文件服务和 AI 服务，再引入网关、注册中心和配置中心。")


def build_report() -> Path:
    ensure_dirs()
    diagrams = make_diagrams()
    screens = preprocess_screens()

    doc = Document()
    style_doc(doc)
    add_cover(doc)
    add_toc(doc)
    add_intro(doc, diagrams, screens)
    add_requirement_chapter(doc, diagrams)
    add_design_chapter(doc, diagrams)
    add_implementation_chapter(doc, screens)
    add_test_chapter(doc)
    add_summary(doc)
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)

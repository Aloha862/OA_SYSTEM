from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PENDING_DIR = ROOT / "课设要求与说明" / "待更新"
OUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = OUT_DIR / "assets_updated"
VERSION = "V1.1"
UPDATE_DATE = "2026年6月9日"

PROJECT_NAME = "企业 OA 管理系统"
PROJECT_ID = "OA_SYSTEM"


COLORS = {
    "ink": (28, 37, 54),
    "muted": (92, 106, 128),
    "line": (56, 73, 98),
    "blue": (219, 235, 255),
    "green": (218, 246, 226),
    "yellow": (255, 244, 207),
    "purple": (238, 228, 255),
    "red": (255, 224, 224),
    "cyan": (216, 245, 247),
    "gray": (247, 249, 252),
    "white": (255, 255, 255),
    "accent": (54, 119, 246),
    "success": (25, 156, 88),
    "warning": (210, 139, 28),
}


MODULES = [
    {
        "name": "用户登录与认证权限",
        "summary": "登录、注册、退出、Token 刷新、当前用户信息、JWT 鉴权、角色与审批人身份控制。",
        "frontend": "login、register、stores/user、utils/request",
        "backend": "AuthController、JwtAuthenticationFilter、SecurityConfig、SecurityPermission",
        "tables": "sys_user",
        "apis": "POST /api/auth/login；POST /api/auth/register；POST /api/auth/logout；GET /api/auth/me；POST /api/auth/refresh",
        "status": "已完成",
    },
    {
        "name": "用户管理",
        "summary": "管理员维护员工账号，支持分页筛选、新增、修改、逻辑删除、批量删除、启停、重置密码、设置审批人；员工维护个人资料。",
        "frontend": "views/user、views/profile、api/users",
        "backend": "UserController、UserServiceImpl、UserMapper",
        "tables": "sys_user",
        "apis": "GET /api/users/page；POST /api/users；PUT /api/users/{id}；DELETE /api/users/{id}；PUT /api/users/{id}/status",
        "status": "已完成",
    },
    {
        "name": "部门组织管理",
        "summary": "维护部门树、部门分页、负责人、默认审批人；审批提交时按部门解析审批人。",
        "frontend": "views/department、api/departments",
        "backend": "DepartmentController、DepartmentServiceImpl、DepartmentMapper",
        "tables": "sys_department、sys_user.department_id",
        "apis": "GET /api/departments/tree；GET /api/departments/page；POST /api/departments；PUT /api/departments/{id}/approver",
        "status": "已完成",
    },
    {
        "name": "审批流程管理",
        "summary": "支持请假、报销、加班、出差四类审批，覆盖草稿、提交、撤回、通过、驳回、记录、AI 摘要和风险提示。",
        "frontend": "views/approval/list、create、todo、detail",
        "backend": "ApprovalController、ApprovalServiceImpl、ApprovalRecordMapper",
        "tables": "oa_approval、oa_approval_record",
        "apis": "GET /api/approvals/page；POST /api/approvals；POST /api/approvals/{id}/submit；POST /api/approvals/{id}/approve；POST /api/approvals/{id}/reject",
        "status": "已完成",
    },
    {
        "name": "日程会议管理",
        "summary": "支持个人日程、会议日程、参与人管理、接受/拒绝、今日/本周/日历视图、定时提醒和自然语言 AI 解析。",
        "frontend": "views/schedule/list、calendar、api/schedules",
        "backend": "ScheduleController、ScheduleServiceImpl、ScheduleReminderTask",
        "tables": "oa_schedule、oa_schedule_participant",
        "apis": "GET /api/schedules/page；GET /api/schedules/calendar；POST /api/schedules；POST /api/schedules/{id}/participants；POST /api/schedules/ai-parse",
        "status": "已完成",
    },
    {
        "name": "新闻公告管理",
        "summary": "管理员维护新闻公告草稿、发布、下架、置顶；员工查看、评论、点赞、收藏；支持 AI 生成和润色。",
        "frontend": "views/news/list、detail、manage、api/news",
        "backend": "NewsController、NewsServiceImpl、NewsCommentMapper、NewsLikeMapper、NewsFavoriteMapper",
        "tables": "oa_news、oa_news_comment、oa_news_like、oa_news_favorite",
        "apis": "GET /api/news/page；POST /api/news；POST /api/news/{id}/publish；POST /api/news/{id}/comments；POST /api/news/{id}/like；POST /api/news/ai-generate",
        "status": "已完成",
    },
    {
        "name": "字典与基础数据",
        "summary": "维护审批类型、审批状态、日程类型、新闻状态、通知类型等枚举数据，提供按类型读取和缓存刷新。",
        "frontend": "views/dict、components/DictSelect、stores/dict",
        "backend": "DictTypeController、DictDataController、DictDataServiceImpl",
        "tables": "sys_dict_type、sys_dict_data",
        "apis": "GET /api/dict-types/page；GET /api/dict-data/type/{typeCode}；POST /api/dict-data/cache/refresh",
        "status": "已完成",
    },
    {
        "name": "文件上传与附件管理",
        "summary": "支持头像、审批附件、日程附件、新闻封面等文件上传、下载、业务关联、列表查询和逻辑删除。",
        "frontend": "components/OaUpload、api/files",
        "backend": "OaFileController、OaFileServiceImpl、FileProperties、WebMvcConfig",
        "tables": "sys_file",
        "apis": "POST /api/files/upload；GET /api/files/{id}；GET /api/files/business；GET /api/files/download/{id}；DELETE /api/files/{id}",
        "status": "已完成",
    },
    {
        "name": "通知消息与实时推送",
        "summary": "站内信分页、未读数、已读、批量已读、删除、系统通知，并串联 RabbitMQ、WebSocket 和邮件提醒。",
        "frontend": "views/notification、NotificationBell、stores/notification、utils/ws",
        "backend": "NotificationController、NotificationProducer、NotificationConsumer、NotificationWebSocketHandler、MailServiceImpl",
        "tables": "sys_notification",
        "apis": "GET /api/notifications/page；GET /api/notifications/unread-count；PUT /api/notifications/{id}/read；POST /api/notifications/system；WebSocket /ws/notification",
        "status": "已完成",
    },
    {
        "name": "AI 助手与调用日志",
        "summary": "提供审批摘要、审批风险、新闻生成、新闻润色、日程解析、智能问答，并记录 AI 调用日志；支持 Mock 与通义千问。",
        "frontend": "views/ai、components/AiDialog、api/ai",
        "backend": "AiController、MockAiServiceImpl、TongyiAiServiceImpl、AiLogServiceImpl",
        "tables": "sys_ai_log",
        "apis": "POST /api/ai/approval-summary；POST /api/ai/approval-risk；POST /api/ai/news-generate；POST /api/ai/schedule-parse；GET /api/ai/logs/page",
        "status": "已完成",
    },
]

TABLES = [
    ("sys_user", "用户表", "账号、密码、姓名、角色、状态、部门、岗位、审批人标识。"),
    ("sys_department", "部门表", "部门名称、上级部门、负责人、默认审批人、排序、状态。"),
    ("oa_approval", "审批主表", "审批编号、标题、类型、状态、申请人、审批人、时间、金额、AI 摘要和风险。"),
    ("oa_approval_record", "审批记录表", "审批单、操作人、动作、意见。"),
    ("oa_schedule", "日程表", "标题、内容、类型、创建人、起止时间、地点、提醒时间和状态。"),
    ("oa_schedule_participant", "日程参与人表", "日程、参与用户、接受/拒绝状态。"),
    ("oa_news", "新闻公告表", "标题、摘要、正文、分类、封面、发布状态、置顶、阅读/点赞/收藏/评论计数。"),
    ("oa_news_comment", "新闻评论表", "新闻、评论用户、内容、父评论、状态。"),
    ("oa_news_like", "新闻点赞表", "新闻与用户的点赞关系。"),
    ("oa_news_favorite", "新闻收藏表", "新闻与用户的收藏关系。"),
    ("sys_dict_type", "字典类型表", "类型编码、类型名称、状态。"),
    ("sys_dict_data", "字典数据表", "类型编码、标签、值、排序、状态。"),
    ("sys_file", "文件表", "原始文件名、存储文件名、路径、URL、业务类型、上传人、下载次数。"),
    ("sys_notification", "站内信通知表", "接收人、发送人、标题、内容、类型、业务 ID、已读状态、推送状态。"),
    ("sys_ai_log", "AI 调用日志表", "用户、功能类型、模型提供方、请求内容、响应内容、耗时、成功状态。"),
]

TECH_STACK = [
    ("后端语言与框架", "Java 17、Spring Boot 3.2.5、Maven"),
    ("安全认证", "Spring Security、JWT、BCrypt、Redis Token 黑名单"),
    ("数据访问", "MySQL 8、MyBatis-Plus 3.5.7、逻辑删除、分页插件、自动填充"),
    ("缓存与消息", "Redis、RabbitMQ TopicExchange、手动 ACK"),
    ("实时通信", "WebSocket，地址 /ws/notification"),
    ("文件与邮件", "本地磁盘上传、/files 静态访问、Spring Mail"),
    ("接口文档", "Springdoc、Knife4j"),
    ("前端", "Vue 3、Vite 6、TypeScript、Element Plus、Pinia、Vue Router、Axios、ECharts、vue-i18n"),
    ("基础设施", "Docker Compose、MySQL、Redis、RabbitMQ management"),
]

FRONTEND_PAGES = [
    ("登录/注册", "/login、/register", "账号登录、员工注册、Token 写入、本地语言切换。"),
    ("首页仪表盘", "/dashboard", "统计概览、近期审批、今日日程、最新新闻。"),
    ("用户管理", "/users", "管理员维护员工账号、状态、角色、审批人标识。"),
    ("部门管理", "/departments", "组织树、负责人、默认审批人维护。"),
    ("审批管理", "/approvals、/approvals/create、/approvals/todo、/approvals/:id", "发起、查询、待办、审批详情和记录。"),
    ("日程管理", "/schedules、/schedules/calendar", "列表视图、日历视图、会议参与和提醒。"),
    ("新闻中心", "/news、/news/:id、/news-manage", "新闻阅读、详情互动、后台发布管理。"),
    ("字典管理", "/dicts", "字典类型与字典项维护。"),
    ("站内信中心", "/notifications", "未读、已读、批量处理和系统通知。"),
    ("AI 助手", "/ai", "智能问答、日程解析，业务页面复用 AI 对话组件。"),
    ("个人中心", "/profile", "头像、联系方式和基础资料维护。"),
]

STATUS_TABLE = [
    ("源码结构", "backend、frontend、docs、scripts、docker-compose.yml 已形成完整工程目录。"),
    ("后端服务", "Spring Boot API 已覆盖 9 个业务模块，统一 Result 返回、统一异常处理、JWT 鉴权和 MyBatis-Plus CRUD。"),
    ("前端页面", "Vue 3 管理端已覆盖登录、仪表盘、用户、部门、审批、日程、新闻、字典、通知、AI、个人中心等页面。"),
    ("数据模型", "schema.sql 覆盖 15 张核心表，所有业务表具备 created_at、updated_at、deleted 逻辑删除字段。"),
    ("异步与实时", "审批、日程、新闻和系统通知可进入 RabbitMQ，消费者落库后通过 WebSocket 推送，并可发送邮件提醒。"),
    ("AI 能力", "MockAiServiceImpl 可用于演示，TongyiAiServiceImpl 预留通义千问兼容接口并记录 sys_ai_log。"),
    ("测试脚本", "scripts/qa-api-regression.ps1 覆盖核心 API 回归流程，可作为联调验收脚本。"),
]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for block in text.split("\n"):
        current = ""
        for char in block:
            trial = current + char
            if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = char
        lines.append(current)
    return [line for line in lines if line]


def draw_centered_text(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str, size: int = 30, color=None, bold=False) -> None:
    color = color or COLORS["ink"]
    fnt = font(size, bold)
    x1, y1, x2, y2 = rect
    lines = wrap_text(draw, text, fnt, x2 - x1 - 36)
    line_height = size + 10
    total_height = line_height * len(lines)
    y = y1 + (y2 - y1 - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=fnt, fill=color)
        y += line_height


def draw_box(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str, fill, outline=None, size=28, radius=18, bold=False) -> None:
    outline = outline or COLORS["line"]
    draw.rounded_rectangle(rect, radius=radius, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, rect, text, size=size, bold=bold)


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str | None = None) -> None:
    draw.text((70, 46), title, font=font(46, True), fill=COLORS["ink"])
    if subtitle:
        draw.text((72, 104), subtitle, font=font(24), fill=COLORS["muted"])
    draw.line((70, 142, 1730, 142), fill=(190, 201, 216), width=3)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=None, width=5) -> None:
    color = color or COLORS["line"]
    draw.line((start, end), fill=color, width=width)
    sx, sy = start
    ex, ey = end
    dx, dy = ex - sx, ey - sy
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 18
    p1 = (ex, ey)
    p2 = (ex - ux * size + px * size * 0.55, ey - uy * size + py * size * 0.55)
    p3 = (ex - ux * size - px * size * 0.55, ey - uy * size - py * size * 0.55)
    draw.polygon([p1, p2, p3], fill=color)


def new_canvas() -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1800, 1050), COLORS["white"])
    return image, ImageDraw.Draw(image)


def save_image(image: Image.Image, name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / name
    image.save(path)
    return path


def create_architecture() -> Path:
    image, draw = new_canvas()
    draw_title(draw, "OA_SYSTEM 系统总体架构图", "前后端分离 + 模块化 Spring Boot + 数据缓存消息与实时推送")
    draw_box(draw, (90, 230, 410, 420), "Vue 3 管理端\nElement Plus\nPinia / Router / Axios", COLORS["blue"], size=26, bold=True)
    draw_box(draw, (650, 210, 1160, 455), "Spring Boot 3 后端\nController / Service / Mapper\n统一返回 / 异常处理 / JWT 鉴权", COLORS["green"], size=28, bold=True)
    draw_box(draw, (1380, 230, 1710, 420), "AI 接入\nMock / 通义千问\nsys_ai_log", COLORS["yellow"], size=28, bold=True)
    draw_box(draw, (120, 690, 430, 860), "MySQL 8\n15 张核心业务表\noa_management", COLORS["red"], size=26)
    draw_box(draw, (535, 690, 845, 860), "Redis\n字典 / 部门 / 新闻\n未读数 / Token 黑名单", COLORS["cyan"], size=24)
    draw_box(draw, (950, 690, 1260, 860), "RabbitMQ\nTopicExchange\n通知队列 / 手动 ACK", COLORS["purple"], size=24)
    draw_box(draw, (1365, 690, 1675, 860), "本地文件目录\nbackend/upload\n/files 静态访问", COLORS["green"], size=24)
    arrow(draw, (410, 305), (650, 305), COLORS["accent"])
    draw.text((455, 262), "REST API / JWT", font=font(24), fill=COLORS["accent"])
    arrow(draw, (650, 380), (410, 380), (111, 77, 230))
    draw.text((440, 392), "WebSocket 通知", font=font(23), fill=(111, 77, 230))
    arrow(draw, (1160, 310), (1380, 310), COLORS["warning"])
    arrow(draw, (815, 455), (275, 690), COLORS["line"], width=3)
    arrow(draw, (855, 455), (690, 690), COLORS["line"], width=3)
    arrow(draw, (900, 455), (1105, 690), COLORS["line"], width=3)
    arrow(draw, (945, 455), (1520, 690), COLORS["line"], width=3)
    return save_image(image, "architecture_v2.png")


def create_modules() -> Path:
    image, draw = new_canvas()
    draw_title(draw, "OA_SYSTEM 功能模块图", "根据当前前端页面、后端 Controller 与数据库表核对生成")
    groups = [
        ("基础支撑", "认证权限\n用户管理\n部门管理\n字典配置", COLORS["blue"]),
        ("核心业务", "审批流程\n日程会议\n新闻公告", COLORS["green"]),
        ("协作能力", "通知推送\n文件附件\nWebSocket / 邮件", COLORS["cyan"]),
        ("智能增强", "AI 摘要\n风险分析\n新闻生成 / 日程解析", COLORS["yellow"]),
    ]
    group_rects = [
        (90, 220, 465, 455),
        (515, 220, 890, 455),
        (940, 220, 1315, 455),
        (1365, 220, 1740, 455),
    ]
    for (title, body, fill), rect in zip(groups, group_rects):
        draw_box(draw, rect, f"{title}\n{body}", fill, size=26, bold=True)
    boxes = [
        ("登录 / 注册 / JWT", COLORS["blue"]),
        ("员工资料 / 启停 / 审批人", COLORS["green"]),
        ("组织树 / 负责人 / 默认审批人", COLORS["cyan"]),
        ("请假 / 报销 / 加班 / 出差", COLORS["yellow"]),
        ("日历视图 / 参与人 / 提醒", COLORS["purple"]),
        ("发布 / 评论 / 点赞 / 收藏", COLORS["red"]),
        ("基础枚举 / 缓存刷新", COLORS["blue"]),
        ("上传 / 下载 / 业务关联", COLORS["green"]),
        ("站内信 / WebSocket / 邮件", COLORS["cyan"]),
        ("摘要 / 风险 / 生成 / 问答", COLORS["yellow"]),
    ]
    start_x, start_y = 90, 570
    box_w, box_h = 300, 115
    x_gap, y_gap = 45, 65
    for idx, (text, fill) in enumerate(boxes):
        row = idx // 5
        col = idx % 5
        x = start_x + col * (box_w + x_gap)
        y = start_y + row * (box_h + y_gap)
        draw_box(draw, (x, y, x + box_w, y + box_h), text, fill, size=22)
    draw.text((90, 945), "说明：上方按能力域划分，下方列出当前已实现的主要功能点。", font=font(24), fill=COLORS["muted"])
    return save_image(image, "modules_v2.png")


def create_database_er() -> Path:
    image, draw = new_canvas()
    draw_title(draw, "OA_SYSTEM 数据库 ER 摘要图", "schema.sql 当前定义 15 张核心表，业务关系主要在 Service 层控制")
    domains = [
        ("基础组织域", (80, 205, 510, 430), COLORS["blue"], ["sys_user", "sys_department"]),
        ("审批域", (555, 205, 985, 430), COLORS["yellow"], ["oa_approval", "oa_approval_record"]),
        ("日程域", (1030, 205, 1460, 430), COLORS["green"], ["oa_schedule", "oa_schedule_participant"]),
        ("新闻内容域", (80, 520, 510, 825), COLORS["red"], ["oa_news", "oa_news_comment", "oa_news_like", "oa_news_favorite"]),
        ("基础支撑域", (555, 520, 985, 825), COLORS["cyan"], ["sys_dict_type", "sys_dict_data", "sys_file"]),
        ("协同与智能域", (1030, 520, 1460, 825), COLORS["purple"], ["sys_notification", "sys_ai_log"]),
    ]
    for title, rect, fill, tables in domains:
        draw.rounded_rectangle(rect, radius=18, fill=fill, outline=COLORS["line"], width=3)
        draw.text((rect[0] + 24, rect[1] + 22), title, font=font(28, True), fill=COLORS["ink"])
        y = rect[1] + 75
        for table_name in tables:
            draw_box(draw, (rect[0] + 26, y, rect[2] - 26, y + 48), table_name, COLORS["white"], outline=(132, 146, 166), size=21)
            y += 58
    relation_panel = (1510, 205, 1735, 825)
    draw.rounded_rectangle(relation_panel, radius=18, fill=COLORS["gray"], outline=COLORS["line"], width=3)
    draw.text((1532, 230), "主要关系", font=font(28, True), fill=COLORS["ink"])
    rel_text = [
        "用户 -> 部门",
        "用户 -> 审批申请/审批人",
        "审批 -> 审批记录",
        "日程 -> 参与人",
        "新闻 -> 评论/点赞/收藏",
        "业务 -> 文件附件",
        "业务事件 -> 站内信",
        "用户 -> AI 日志",
        "字典类型 -> 字典数据",
    ]
    y = 290
    for text in rel_text:
        draw.text((1532, y), "• " + text, font=font(19), fill=COLORS["ink"])
        y += 45
    arrow(draw, (510, 318), (555, 318), COLORS["line"], width=4)
    arrow(draw, (985, 318), (1030, 318), COLORS["line"], width=4)
    arrow(draw, (510, 632), (555, 632), COLORS["line"], width=4)
    arrow(draw, (985, 632), (1030, 632), COLORS["line"], width=4)
    draw.text((90, 905), "说明：当前 schema.sql 未显式声明数据库外键，避免课程设计联调阶段因数据顺序造成约束阻塞；业务一致性由 Service 校验和唯一索引控制。", font=font(24), fill=COLORS["muted"])
    return save_image(image, "database_er_v2.png")


def create_linear_flow(title: str, subtitle: str, steps: Sequence[str], color, name: str) -> Path:
    image, draw = new_canvas()
    draw_title(draw, title, subtitle)
    y1, y2 = 430, 575
    gap = 32
    box_w = int((1660 - gap * (len(steps) - 1)) / len(steps))
    x = 70
    for idx, step in enumerate(steps):
        rect = (x, y1, x + box_w, y2)
        draw_box(draw, rect, step, COLORS["gray"], outline=color, size=23)
        if idx < len(steps) - 1:
            arrow(draw, (x + box_w, (y1 + y2) // 2), (x + box_w + gap, (y1 + y2) // 2), color, width=5)
        x += box_w + gap
    return save_image(image, name)


def create_approval_flow() -> Path:
    return create_linear_flow(
        "审批业务流程图",
        "单级自研审批流：草稿提交后自动匹配部门审批人，并通过通知链路反馈结果",
        ["创建审批草稿", "提交审批\n生成编号", "匹配部门审批人", "写入审批记录\n发送通知", "审批人处理", "更新状态\n通知申请人"],
        COLORS["success"],
        "approval_flow_v2.png",
    )


def create_notification_flow() -> Path:
    return create_linear_flow(
        "通知消息流转图",
        "业务事件先进入 RabbitMQ，消费者负责落库、WebSocket 在线推送和邮件提醒",
        ["业务事件", "NotificationProducer", "RabbitMQ\nTopicExchange", "NotificationConsumer", "站内信落库", "WebSocket 推送\n邮件提醒"],
        (111, 77, 230),
        "notification_flow_v2.png",
    )


def create_login_flow() -> Path:
    return create_linear_flow(
        "登录认证与权限流程图",
        "JWT 作为主要登录凭证，前端拦截器携带 Token，后端过滤器解析并写入 SecurityContext",
        ["输入账号密码", "BCrypt 校验", "签发 JWT", "前端保存 Token", "请求携带\nAuthorization", "后端鉴权\n角色/审批人判断"],
        COLORS["accent"],
        "login_flow_v2.png",
    )


def create_deployment() -> Path:
    image, draw = new_canvas()
    draw_title(draw, "OA_SYSTEM 部署结构图", "本地课程设计环境：浏览器 + 前端构建产物 + Spring Boot + MySQL/Redis/RabbitMQ")
    draw_box(draw, (100, 240, 410, 410), "浏览器用户\n管理员 / 员工", COLORS["blue"], size=27)
    draw_box(draw, (565, 240, 875, 410), "Vue 3 前端\nVite 开发或 dist 部署", COLORS["green"], size=26)
    draw_box(draw, (1030, 240, 1340, 410), "Spring Boot 后端\nlocalhost:8081", COLORS["yellow"], size=26)
    draw_box(draw, (100, 690, 390, 835), "MySQL 8\n业务数据", COLORS["red"], size=25)
    draw_box(draw, (485, 690, 775, 835), "Redis 7\n缓存 / 黑名单", COLORS["cyan"], size=25)
    draw_box(draw, (870, 690, 1160, 835), "RabbitMQ\n通知队列", COLORS["purple"], size=25)
    draw_box(draw, (1255, 690, 1545, 835), "本地上传目录\nbackend/upload", COLORS["green"], size=25)
    draw_box(draw, (1390, 240, 1700, 410), "外部 AI 服务\nDashScope / Qwen", COLORS["yellow"], size=25)
    arrow(draw, (410, 325), (565, 325), COLORS["accent"])
    arrow(draw, (875, 325), (1030, 325), COLORS["accent"])
    arrow(draw, (1340, 325), (1390, 325), COLORS["warning"])
    for x in [245, 630, 1015, 1400]:
        arrow(draw, (1185, 410), (x, 690), COLORS["line"], width=3)
    draw.text((1050, 475), "JDBC / RedisTemplate / RabbitTemplate / Multipart", font=font(23), fill=COLORS["muted"])
    return save_image(image, "deployment_v2.png")


def create_use_case() -> Path:
    image, draw = new_canvas()
    draw_title(draw, "OA_SYSTEM 用例图", "管理员、普通员工与审批人三类使用视角")
    lanes = [
        ("管理员\nADMIN", COLORS["blue"], ["用户/部门/字典维护", "新闻发布/置顶/下架", "系统通知/AI 日志"]),
        ("员工\nEMPLOYEE", COLORS["green"], ["发起审批申请", "日程会议管理", "新闻查看/评论/点赞收藏"]),
        ("审批人\nis_approver", COLORS["yellow"], ["处理待办审批", "查看审批记录", "使用审批 AI 摘要/风险"]),
    ]
    y = 210
    for actor, fill, cases in lanes:
        draw_box(draw, (90, y, 365, y + 160), actor, fill, size=27, bold=True)
        x = 500
        for case in cases:
            draw_box(draw, (x, y + 25, x + 330, y + 135), case, COLORS["gray"], outline=COLORS["line"], size=23)
            arrow(draw, (365, y + 80), (500, y + 80), COLORS["line"], width=3)
            x += 390
        y += 245
    draw_box(draw, (500, 890, 1610, 990), "公共用例：登录认证、个人资料、站内信、基础 AI 问答、文件附件访问", COLORS["purple"], size=25)
    return save_image(image, "use_case_v2.png")


def generate_images() -> dict[str, Path]:
    return {
        "architecture": create_architecture(),
        "modules": create_modules(),
        "database": create_database_er(),
        "approval_flow": create_approval_flow(),
        "notification_flow": create_notification_flow(),
        "login_flow": create_login_flow(),
        "deployment": create_deployment(),
        "use_case": create_use_case(),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, True)
        set_cell_shading(hdr[idx], "D9EAF7")
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(90, 90, 90)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_paragraph(doc: Document, text: str = "", style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
    p.paragraph_format.first_line_indent = Cm(0.74) if style is None and text else None
    p.paragraph_format.line_spacing = 1.35


def add_bullets(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(31, 45, 61)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "黑体"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        styles[style_name].font.color.rgb = RGBColor(31, 45, 61)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def add_cover(doc: Document, en_title: str, cn_title: str, doc_type: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(en_title)
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    run.bold = True
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cn_title)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(24)
    run.bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_type)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(20)
    run.bold = True
    for _ in range(4):
        doc.add_paragraph()
    add_table(
        doc,
        ["项目", "内容"],
        [
            ("项目名称", f"{PROJECT_NAME}（{PROJECT_ID}）"),
            ("文档版本", VERSION),
            ("更新日期", UPDATE_DATE),
            ("更新依据", "当前 backend、frontend、db、scripts 目录源码及实际实现情况"),
            ("班级/小组", "提交前由小组填写"),
            ("指导教师", "提交前由小组填写"),
        ],
        [4.0, 10.5],
    )
    add_paragraph(doc, "说明：本文档已按当前项目实际完成情况更新，保留班级、小组、姓名等需人工确认的信息占位。")
    doc.add_page_break()


def add_toc(doc: Document, items: Sequence[tuple[int, str]]) -> None:
    add_heading(doc, "目  录", 1)
    for level, text in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm((level - 1) * 0.65)
        run = p.add_run(text)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
    doc.add_page_break()


def add_common_status(doc: Document) -> None:
    add_table(doc, ["检查项", "当前完成情况"], STATUS_TABLE, [4.0, 11.0])


def add_tech_stack(doc: Document) -> None:
    add_table(doc, ["类别", "当前采用技术"], TECH_STACK, [4.0, 11.0])


def add_module_overview_table(doc: Document) -> None:
    add_table(
        doc,
        ["模块", "主要能力", "前端/后端/数据表", "状态"],
        [
            (
                m["name"],
                m["summary"],
                f"前端：{m['frontend']}\n后端：{m['backend']}\n数据：{m['tables']}",
                m["status"],
            )
            for m in MODULES
        ],
        [3.2, 6.0, 5.2, 1.5],
    )


def add_boundary_notes(doc: Document) -> None:
    add_bullets(
        doc,
        [
            "当前后端是单体模块化 Spring Boot 应用，不是多个独立部署微服务；文档中按课程要求说明架构分层和可演进方向。",
            "审批流程采用自研单级审批模型，未接入 Activiti、Flowable 或 BPMN 引擎。",
            "文件存储采用本地目录 backend/upload，并通过 /files 静态路径访问，未接入 OSS 或对象存储。",
            "AI 能力包含 Mock 与通义千问兼容接口；正式调用依赖 DASHSCOPE_API_KEY 或 TONGYI_API_KEY。",
            "结构化日志当前主要落在 sys_ai_log；系统运行日志由 Spring Boot 日志输出承担。",
        ]
    )


def add_requirements_module_details(doc: Document) -> None:
    for idx, module in enumerate(MODULES, start=1):
        add_heading(doc, f"3.{idx} {module['name']}", 2)
        add_paragraph(doc, module["summary"])
        add_table(
            doc,
            ["需求项", "说明"],
            [
                ("业务目标", module["summary"]),
                ("主要接口", module["apis"]),
                ("前端页面/组件", module["frontend"]),
                ("后端实现类", module["backend"]),
                ("涉及数据表", module["tables"]),
                ("完成状态", module["status"]),
            ],
            [3.3, 11.8],
        )


def add_design_module_details(doc: Document) -> None:
    details = [
        (
            "用户登录与认证权限模块",
            "AuthController 接收登录、注册、退出和刷新 Token 请求；UserServiceImpl 校验 BCrypt 密码、签发 JWT、写入最后登录时间；JwtAuthenticationFilter 在每次请求中解析 Authorization 请求头，检查 Redis Token 黑名单，并把 LoginUser 写入 SecurityContext；SecurityConfig 配置匿名路径和接口授权规则。",
        ),
        (
            "用户管理模块",
            "UserController 面向管理员提供分页、详情、新增、修改、删除、批量删除、启停、重置密码和审批人身份维护接口，面向登录用户提供个人资料查看和修改接口。UserServiceImpl 使用 MyBatis-Plus LambdaQueryWrapper 组合筛选条件，新增和重置密码统一使用 BCrypt。",
        ),
        (
            "部门组织管理模块",
            "DepartmentServiceImpl 负责部门树构造、分页查询、负责人维护和默认审批人维护。部门树使用 Redis 缓存，部门变更后清理缓存；审批提交时通过部门默认审批人或负责人解析审批人。",
        ),
        (
            "审批流程管理模块",
            "ApprovalServiceImpl 维护 DRAFT、PENDING、APPROVED、REJECTED、WITHDRAWN 状态流转。申请人创建草稿后提交，系统生成审批编号并匹配审批人；审批人处理时写入 oa_approval_record，并通过 NotificationProducer 触发通知。",
        ),
        (
            "日程会议管理模块",
            "ScheduleServiceImpl 维护个人日程和会议日程，保存会议参与人并控制参与人接受/拒绝状态。ScheduleReminderTask 每分钟扫描到达提醒时间的日程，投递提醒事件，消费者落库并推送站内信。",
        ),
        (
            "新闻公告管理模块",
            "NewsServiceImpl 支持草稿、发布、下架、置顶、评论、点赞和收藏；新闻详情使用 Redis 缓存，修改、删除、发布和下架时清理缓存；发布和评论事件可触发通知。",
        ),
        (
            "字典与基础数据模块",
            "DictTypeServiceImpl 维护字典类型，DictDataServiceImpl 维护字典项。前端表单通过 typeCode 获取选项，后端优先读取 Redis 缓存，字典变更后清理或刷新缓存。",
        ),
        (
            "文件上传与附件管理模块",
            "OaFileServiceImpl 接收 MultipartFile，按年月目录保存到本地 upload 目录，使用 UUID 重命名，记录原始文件名、访问 URL、业务类型、业务 ID 和上传人。WebMvcConfig 映射 /files 静态资源路径。",
        ),
        (
            "通知消息与实时推送模块",
            "业务服务构造 NotificationMessage 后通过 RabbitMQ TopicExchange 投递；NotificationConsumer 消费消息后写入 sys_notification，并调用 NotificationWebSocketHandler 给在线用户推送，必要时调用 MailServiceImpl 发送邮件提醒。",
        ),
        (
            "AI 助手与调用日志模块",
            "AiController 暴露审批摘要、审批风险、新闻生成、新闻润色、日程解析和智能问答接口。MockAiServiceImpl 用于无密钥演示，TongyiAiServiceImpl 通过 OpenAI 兼容接口调用通义千问，AiLogServiceImpl 记录请求、响应、模型和耗时。",
        ),
    ]
    for idx, (title, body) in enumerate(details, start=1):
        add_heading(doc, f"3.{idx} {title}", 2)
        add_paragraph(doc, body)
        module = MODULES[idx - 1]
        add_table(
            doc,
            ["设计项", "内容"],
            [
                ("核心类", module["backend"]),
                ("接口路径", module["apis"]),
                ("数据表", module["tables"]),
                ("前端入口", module["frontend"]),
            ],
            [3.0, 12.0],
        )


def save_doc(doc: Document, filename: str) -> Path:
    PENDING_DIR.mkdir(parents=True, exist_ok=True)
    path = PENDING_DIR / filename
    doc.save(path)
    return path


def build_proposal(images: dict[str, Path]) -> Path:
    doc = Document()
    configure_doc(doc)
    add_cover(doc, "Project Start Report", f"{PROJECT_NAME}项目立项报告", "项目立项说明书")
    add_toc(
        doc,
        [
            (1, "1 Project Proposal 项目提出"),
            (2, "1.1 Project Brief 项目简介"),
            (2, "1.2 Project Goal 项目目标"),
            (2, "1.3 System Scope 系统边界"),
            (2, "1.4 Estimated Effort 工作量估计"),
            (1, "2 Team building and Schedule 开发团队组成和计划时间"),
            (2, "2.1 Project Team 开发团队"),
            (2, "2.2 Project Plan 计划时间"),
            (1, "3 Evaluating and Mitigating 风险评估和规避"),
            (2, "3.1 Technical Risks 技术风险"),
            (2, "3.2 Management Risks 管理风险"),
            (2, "3.3 Current Completion 当前完成情况"),
        ],
    )
    add_heading(doc, "1 Project Proposal 项目提出", 1)
    add_heading(doc, "1.1 Project Brief 项目简介", 2)
    add_paragraph(doc, f"{PROJECT_ID} 是一个面向企业日常办公场景的前后端分离管理系统。系统围绕用户、部门、审批、日程、新闻、字典、文件、通知和 AI 助手等办公协同场景展开，后端提供统一 REST API，前端提供基于 Vue 3 和 Element Plus 的管理端界面。")
    add_paragraph(doc, "经当前源码核查，项目后端为单体模块化 Spring Boot 应用，具备清晰的模块包结构和可演进为微服务的边界；课程设计报告中可按服务能力分层说明，但不应描述为已经拆分部署的多微服务系统。")
    add_picture(doc, images["modules"], "图 1-1 企业 OA 管理系统功能结构图")

    add_heading(doc, "1.2 Project Goal 项目目标", 2)
    add_bullets(
        doc,
        [
            "实现管理员、普通员工和审批人三类使用视角，完成登录认证、权限控制和个人资料维护。",
            "实现用户、部门、字典、文件等基础管理能力，为审批、日程、新闻和通知模块提供基础数据。",
            "实现审批申请、单级审批、审批记录、附件关联和审批结果通知。",
            "实现日程会议、会议参与人、提醒扫描、RabbitMQ 通知、WebSocket 实时推送和邮件提醒。",
            "实现新闻公告发布、下架、置顶、评论、点赞、收藏和 AI 辅助创作。",
            "实现 AI 助手的 Mock 演示和通义千问接入配置，保留 AI 调用日志。",
        ]
    )
    add_heading(doc, "1.3 System Scope 系统边界", 2)
    add_paragraph(doc, "系统边界包括浏览器端管理界面、Spring Boot 后端 API、MySQL 数据库、Redis 缓存、RabbitMQ 消息队列、WebSocket 通知通道、本地文件目录和外部 AI 服务。系统不包含移动端 App、企业微信/飞书等第三方组织通讯录同步，也未接入对象存储和 BPMN 流程引擎。")
    add_picture(doc, images["architecture"], "图 1-2 系统总体架构图")
    add_heading(doc, "1.4 Estimated Effort 工作量估计", 2)
    add_table(
        doc,
        ["工作包", "主要内容", "估计占比"],
        [
            ("基础工程与数据库", "后端工程、前端工程、Docker Compose、15 张核心表、初始化数据", "20%"),
            ("基础管理模块", "用户、部门、字典、文件上传、权限控制", "25%"),
            ("核心业务模块", "审批、日程、新闻、通知消息、WebSocket 推送", "35%"),
            ("AI 与联调测试", "Mock/通义 AI、日志、回归脚本、文档和截图", "20%"),
        ],
        [4.0, 8.5, 2.3],
    )

    add_heading(doc, "2 Team building and Schedule 开发团队组成和计划时间", 1)
    add_heading(doc, "2.1 Project Team 开发团队", 2)
    add_table(
        doc,
        ["角色", "职责", "成员"],
        [
            ("项目负责人", "需求拆分、进度管理、最终文档整合", "提交前填写"),
            ("后端开发", "Spring Boot、数据库、接口、权限、缓存、消息队列", "提交前填写"),
            ("前端开发", "Vue 页面、路由、状态管理、API 联调、交互优化", "提交前填写"),
            ("测试与文档", "回归测试、缺陷记录、说明书和答辩材料", "提交前填写"),
        ],
        [3.5, 8.0, 3.2],
    )
    add_heading(doc, "2.2 Project Plan 计划时间", 2)
    add_table(
        doc,
        ["阶段", "计划内容", "交付物"],
        [
            ("第 1 周", "工程初始化、数据库设计、统一返回、认证鉴权", "可启动后端、建表 SQL、登录接口"),
            ("第 2 周", "用户、部门、字典、文件和基础页面", "基础管理模块可联调"),
            ("第 3 周", "审批、日程、新闻、通知、WebSocket 和邮件提醒", "核心业务流程可演示"),
            ("第 4 周", "AI 接入、回归测试、文档更新、答辩材料整理", "最终源码、数据库、说明书和报告"),
        ],
        [3.0, 8.0, 4.2],
    )

    add_heading(doc, "3 Evaluating and Mitigating 风险评估和规避", 1)
    add_heading(doc, "3.1 Technical Risks 技术风险", 2)
    add_table(
        doc,
        ["风险", "影响", "规避措施"],
        [
            ("课程题目与实际架构表述不一致", "若描述为已拆分微服务，答辩时容易被追问", "按实际写成单体模块化 Spring Boot，并说明可按模块演进为微服务。"),
            ("Redis/RabbitMQ/MySQL 环境差异", "服务无法启动或通知链路失败", "提交前统一 application.yml 与 docker-compose 端口、账号配置。"),
            ("AI 密钥不可用", "通义千问调用失败", "保留 Mock 服务作为演示兜底，通过配置切换 provider。"),
            ("本地文件存储缺少持久化策略", "重装或迁移后附件丢失", "提交时保留 upload 示例目录，生产环境需挂载持久化磁盘。"),
            ("中文资源编码", "页面或初始化数据可能出现乱码", "统一 UTF-8 保存源码、SQL 和文档，初始化前先抽样检查显示。"),
        ],
        [4.0, 5.0, 6.0],
    )
    add_heading(doc, "3.2 Management Risks 管理风险", 2)
    add_bullets(
        doc,
        [
            "接口变更导致前后端不一致：以 docs/api-overview.md 和 Controller 实际路径为准，每次变更同步文档。",
            "核心模块延期：优先保证登录、用户、部门、审批、通知链路可演示，再扩展新闻、AI 和统计细节。",
            "文档与代码偏离：最终说明书以 backend、frontend、schema.sql、qa-api-regression.ps1 的当前状态为准更新。",
        ]
    )
    add_heading(doc, "3.3 Current Completion 当前完成情况", 2)
    add_common_status(doc)
    return save_doc(doc, "项目立项说明书—程序员.docx")


def build_requirements(images: dict[str, Path]) -> Path:
    doc = Document()
    configure_doc(doc)
    add_cover(doc, "Soft Requirements Specification", f"{PROJECT_NAME}软件需求说明书", "需求说明书")
    add_toc(
        doc,
        [
            (1, "1 Introduction 简介"),
            (2, "1.1 Purpose 文档目的"),
            (2, "1.2 Scope 本文档适用范围"),
            (1, "2 General description 总体概述"),
            (2, "2.1 Soft perspective 软件概述"),
            (2, "2.2 Soft function 软件功能"),
            (2, "2.3 User Roles and Permissions 用户角色与权限"),
            (1, "3 Functional Requirements 功能需求"),
            *[(2, f"3.{idx} {module['name']}") for idx, module in enumerate(MODULES, start=1)],
            (1, "4 Performance Requirements 非功能需求"),
            (2, "4.1 UI Requirements 界面要求"),
            (2, "4.2 Development Environment 开发环境"),
            (2, "4.3 Development Rules 开发规范"),
            (1, "5 Demand Classification 需求分级"),
        ],
    )
    add_heading(doc, "1 Introduction 简介", 1)
    add_heading(doc, "1.1 Purpose 文档目的", 2)
    add_paragraph(doc, "本文档用于明确企业 OA 管理系统的功能需求、非功能需求、用户角色、接口边界和验收依据，作为本项目开发、测试、文档评审和课程设计答辩的需求基线。")
    add_heading(doc, "1.2 Scope 本文档适用范围", 2)
    add_paragraph(doc, "本文档适用于 OA_SYSTEM 项目的前端页面、后端 API、数据库脚本、消息推送、文件上传和 AI 助手功能说明，不用于描述尚未实现的移动端、复杂多级工作流或第三方组织通讯录集成。")

    add_heading(doc, "2 General description 总体概述", 1)
    add_heading(doc, "2.1 Soft perspective 软件概述", 2)
    add_paragraph(doc, "系统面向企业内部办公自动化场景，目标用户包括管理员、普通员工和具备审批身份的员工。管理员维护基础数据和全局业务，员工处理个人办公事项，审批人处理分配到自己的审批待办。")
    add_picture(doc, images["use_case"], "图 2-1 企业 OA 管理系统用例图")
    add_heading(doc, "2.2 Soft function 软件功能", 2)
    add_picture(doc, images["modules"], "图 2-2 软件功能结构图")
    add_module_overview_table(doc)
    add_heading(doc, "2.3 User Roles and Permissions 用户角色与权限", 2)
    add_table(
        doc,
        ["角色", "权限范围"],
        [
            ("管理员 ADMIN", "可维护用户、部门、字典、文件、新闻公告、系统通知，可查看和管理全局审批、日程、AI 日志。"),
            ("普通员工 EMPLOYEE", "可登录系统、维护个人资料、发起审批、查看本人审批和日程、查看新闻、评论、点赞、收藏、查看站内信、使用基础 AI 助手。"),
            ("审批人 Approver", "由 ADMIN 角色或 sys_user.is_approver=1 标识，除员工能力外，可查看并处理分配到自己的审批待办，可使用审批摘要和风险提示。"),
        ],
        [4.0, 11.0],
    )

    add_heading(doc, "3 Functional Requirements 功能需求", 1)
    add_requirements_module_details(doc)
    add_picture(doc, images["approval_flow"], "图 3-1 审批业务流程图")
    add_picture(doc, images["notification_flow"], "图 3-2 通知消息流转图")

    add_heading(doc, "4 Performance Requirements 非功能需求", 1)
    add_heading(doc, "4.1 UI Requirements 界面要求", 2)
    add_bullets(
        doc,
        [
            "前端采用后台管理系统布局，包含侧边菜单、顶部导航、面包屑、通知入口和头像菜单。",
            "列表页支持分页、筛选、刷新和危险操作二次确认；表单页提供必填校验、提交 loading 和成功/失败提示。",
            "页面至少适配 1366px 宽度，主要表格、弹窗和状态标签保持统一风格。",
            "前端已提供中文和英文资源结构，提交前需统一检查中文资源编码。"
        ],
    )
    add_heading(doc, "4.2 Development Environment 开发环境", 2)
    add_tech_stack(doc)
    add_heading(doc, "4.3 Development Rules 开发规范", 2)
    add_bullets(
        doc,
        [
            "后端接口统一使用 /api 前缀，返回 Result<T> 结构，业务异常以 JSON 返回。",
            "除登录和注册外，业务接口默认需要 JWT；管理员接口使用 hasRole('ADMIN') 或自定义权限方法控制。",
            "数据库表使用 BIGINT 自增主键，核心业务表统一包含 created_at、updated_at、deleted 字段。",
            "高频读取数据优先使用 Redis 缓存，变更后清理或刷新缓存。",
            "通知类业务优先进入 RabbitMQ，消费者负责落库和 WebSocket 推送，避免主流程被邮件等外部服务阻塞。",
        ]
    )
    add_heading(doc, "5 Demand Classification 需求分级", 1)
    add_table(
        doc,
        ["级别", "需求范围", "当前状态"],
        [
            ("必须完成", "登录认证、用户、部门、审批、日程、新闻、字典、文件、通知、数据库脚本、前端主页面", "已完成"),
            ("重要增强", "Redis 缓存、RabbitMQ 异步通知、WebSocket 实时推送、邮件提醒、AI Mock 和通义接入", "已完成/可配置"),
            ("可选扩展", "多级 BPMN 流程、移动端、对象存储、企业通讯录同步、完整审计日志", "未纳入本次实现范围"),
        ],
        [3.0, 9.0, 3.5],
    )
    return save_doc(doc, "需求说明书—程序员.docx")


def build_design(images: dict[str, Path]) -> Path:
    doc = Document()
    configure_doc(doc)
    add_cover(doc, "System Design Specification", f"{PROJECT_NAME}系统设计说明书", "系统设计说明书")
    add_toc(
        doc,
        [
            (1, "1 Introduction 简介"),
            (2, "1.1 Purpose 文档目的"),
            (2, "1.2 Scope 本文档适用范围"),
            (2, "1.3 Name 软件名称"),
            (2, "1.4 Applications 软件应用领域"),
            (1, "2 High Level Design 概要设计"),
            (2, "2.1 系统功能设计"),
            (2, "2.2 系统架构设计"),
            (3, "2.2.1 前后端分离架构"),
            (3, "2.2.2 认证与权限架构"),
            (3, "2.2.3 通知、RabbitMQ 与 WebSocket 链路"),
            (3, "2.2.4 部署结构设计"),
            (2, "2.3 Database 数据库设计"),
            (3, "2.3.1 表关系"),
            (3, "2.3.2 数据表设计"),
            (1, "3 Low Level Model Design 模块详细设计"),
            *[(2, f"3.{idx} {module['name']}") for idx, module in enumerate(MODULES, start=1)],
            (1, "4 Test and Acceptance 测试与验收"),
        ],
    )
    add_heading(doc, "1 Introduction 简介", 1)
    add_heading(doc, "1.1 Purpose 文档目的", 2)
    add_paragraph(doc, "本文档从概要设计和模块详细设计两个层面说明 OA_SYSTEM 的系统结构、数据库、核心类、接口和部署方式，作为代码维护、联调测试和课程验收的设计依据。")
    add_heading(doc, "1.2 Scope 本文档适用范围", 2)
    add_paragraph(doc, "本文档适用于企业 OA 管理系统的后端、前端、数据库、通知链路、文件上传和 AI 助手设计说明。")
    add_heading(doc, "1.3 Name 软件名称", 2)
    add_paragraph(doc, f"软件名称：{PROJECT_NAME}；项目标识：{PROJECT_ID}。")
    add_heading(doc, "1.4 Applications 软件应用领域", 2)
    add_paragraph(doc, "系统应用于企业内部办公管理场景，包括员工资料维护、组织管理、审批流转、会议日程、信息发布、站内通知和 AI 辅助办公。")

    add_heading(doc, "2 High Level Design 概要设计", 1)
    add_heading(doc, "2.1 系统功能设计", 2)
    add_picture(doc, images["modules"], "图 2-1 系统功能设计图")
    add_heading(doc, "2.2 系统架构设计", 2)
    add_picture(doc, images["architecture"], "图 2-2 系统总体架构设计图")
    add_heading(doc, "2.2.1 前后端分离架构", 3)
    add_paragraph(doc, "前端基于 Vue 3 构建 SPA 管理端，后端基于 Spring Boot 提供 REST API，二者通过 Axios 和 JSON 数据交互。前端路由按角色控制菜单显示，后端接口通过 Spring Security 和 JWT 进行访问控制。")
    add_heading(doc, "2.2.2 认证与权限架构", 3)
    add_picture(doc, images["login_flow"], "图 2-3 登录认证与权限流程图")
    add_heading(doc, "2.2.3 通知、RabbitMQ 与 WebSocket 链路", 3)
    add_picture(doc, images["notification_flow"], "图 2-4 通知、消息队列与实时推送链路图")
    add_heading(doc, "2.2.4 部署结构设计", 3)
    add_picture(doc, images["deployment"], "图 2-5 系统部署结构图")
    add_paragraph(doc, "当前 application.yml 中后端端口为 8081，文件上传目录为 E:/code/OA_SYSTEM/backend/upload，AI provider 可配置为 mock 或 tongyi。Docker Compose 提供 MySQL、Redis 和 RabbitMQ 基础依赖，联调前需确保端口和账号配置一致。")

    add_heading(doc, "2.3 Database 数据库设计", 2)
    add_heading(doc, "2.3.1 表关系", 3)
    add_picture(doc, images["database"], "图 2-6 数据库 ER 摘要图")
    add_heading(doc, "2.3.2 数据表设计", 3)
    add_table(doc, ["表名", "中文含义", "主要字段/用途"], TABLES, [4.0, 3.5, 8.0])

    add_heading(doc, "3 Low Level Model Design 模块详细设计", 1)
    add_design_module_details(doc)
    add_picture(doc, images["approval_flow"], "图 3-1 审批模块业务流程图")

    add_heading(doc, "4 Test and Acceptance 测试与验收", 1)
    add_paragraph(doc, "项目提供 scripts/qa-api-regression.ps1 作为接口回归脚本，默认 BaseUrl 为 http://127.0.0.1:8081/api。脚本覆盖管理员、员工和审批人登录，并对部门、字典、新闻、审批、日程、文件、通知和 AI 等主要接口执行创建、查询、修改、状态流转和清理。")
    add_table(
        doc,
        ["验收项", "依据", "期望结果"],
        [
            ("后端启动", "backend/pom.xml、OaApplication", "服务启动，Swagger/Knife4j 可访问。"),
            ("数据库初始化", "schema.sql、data.sql、seed_20_rows.sql", "15 张核心表创建成功，基础用户和字典数据可用。"),
            ("前端构建", "frontend/package.json、src/router/routes.ts", "主要页面可路由访问，构建产物输出到 dist。"),
            ("接口联调", "scripts/qa-api-regression.ps1", "核心 CRUD、审批流转、通知、文件、AI mock/配置接口可通过。"),
            ("文档一致性", "本说明书与源码核查", "文档中的模块、表名、接口路径与当前项目保持一致。"),
        ],
        [3.2, 6.0, 6.0],
    )
    add_heading(doc, "5 Design Boundary 设计边界说明", 1)
    add_boundary_notes(doc)
    return save_doc(doc, "系统设计说明书—程序员.docx")


def main() -> None:
    images = generate_images()
    docs = [build_proposal(images), build_requirements(images), build_design(images)]
    for path in docs:
        print(path)


if __name__ == "__main__":
    main()

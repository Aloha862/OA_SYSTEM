from __future__ import annotations

from pathlib import Path
from typing import Iterable
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"E:\code\OA_SYSTEM")
OUT = ROOT / "output" / "doc"
ASSETS = OUT / "assets"

FONT_HEI = r"C:\Windows\Fonts\simhei.ttf"
FONT_SONG = r"C:\Windows\Fonts\simsun.ttc"

TEAM = [
    ("何成翔", "组长", "00000000000"),
    ("万宇轩", "组员", "00000000000"),
    ("刘灿", "组员", "00000000000"),
    ("姜周良", "组员", "00000000000"),
    ("吴发迟", "组员", "00000000000"),
]

PROJECT_CN = "企业 OA 管理系统"
PROJECT_EN = "OA_SYSTEM"
CLASS_NAME = "23gb软件1班"
TEACHER = "尹菠"


TECH_STACK = [
    ("后端语言与框架", "Java 17、Spring Boot 3.2.5、Maven"),
    ("安全认证", "Spring Security、JWT、BCrypt、Redis Token 黑名单"),
    ("数据访问", "MySQL 8、MyBatis-Plus、逻辑删除、分页插件"),
    ("缓存与消息", "Redis、RabbitMQ TopicExchange、手动 ACK"),
    ("实时通信", "WebSocket，地址 /ws/notification"),
    ("文件与邮件", "本地磁盘上传、Spring Mail"),
    ("接口文档", "Swagger / Knife4j"),
    ("前端", "Vue 3、Vite、TypeScript、Element Plus、Pinia、Vue Router、Axios、ECharts"),
    ("基础设施", "Docker Compose、MySQL、Redis、RabbitMQ management"),
]

TABLES = [
    ("sys_user", "用户表", "登录账号、密码、姓名、角色、部门、岗位、审批人身份、状态"),
    ("sys_department", "部门表", "部门名称、父级部门、负责人、默认审批人、排序、状态"),
    ("oa_approval", "审批表", "审批编号、标题、类型、状态、申请人、审批人、金额、AI 摘要/风险"),
    ("oa_approval_record", "审批记录表", "审批 ID、操作人、动作、审批意见"),
    ("oa_schedule", "日程表", "标题、内容、类型、创建人、起止时间、地点、提醒时间"),
    ("oa_schedule_participant", "日程参与人表", "日程 ID、用户 ID、参与状态"),
    ("oa_news", "新闻表", "标题、摘要、正文、分类、封面、状态、阅读/点赞/收藏/评论数"),
    ("oa_news_comment", "新闻评论表", "新闻 ID、评论用户、内容、父评论、状态"),
    ("oa_news_like", "新闻点赞表", "新闻 ID、用户 ID"),
    ("oa_news_favorite", "新闻收藏表", "新闻 ID、用户 ID"),
    ("sys_dict_type", "字典类型表", "类型编码、类型名称、备注、状态"),
    ("sys_dict_data", "字典数据表", "类型编码、标签、值、排序、状态"),
    ("sys_file", "文件表", "原始文件名、存储文件名、路径、URL、业务类型、上传人"),
    ("sys_notification", "站内信通知表", "接收人、发送人、标题、内容、类型、业务 ID、已读状态"),
    ("sys_ai_log", "AI 调用日志表", "用户、功能类型、模型提供商、请求内容、响应内容、耗时"),
]

MODULES = [
    {
        "title": "用户登录与认证权限模块",
        "functions": "完成登录、注册、登出、Token 刷新、当前用户信息获取，以及管理员和员工角色的接口访问控制。",
        "flow": "用户在登录页提交账号和密码，AuthController 调用 UserService 校验账号状态和 BCrypt 密码，校验通过后 JwtTokenUtil 生成 JWT。前端保存 Token 并在 Axios 请求拦截器中附加 Authorization 请求头。后端 JwtAuthenticationFilter 解析 Token，加载 LoginUser 并写入 SecurityContext，接口再通过 @PreAuthorize 或业务校验限制访问。",
        "classes": "AuthController.login/register/logout/me/refresh；UserServiceImpl.login/logout/refresh/register/me；JwtAuthenticationFilter；JwtTokenUtil；SecurityConfig；CustomUserDetailsService；SecurityPermission。",
        "tables": "sys_user：username、password、role、status、department_id、position、is_approver、last_login_time。",
        "frontend": "frontend/src/views/login/index.vue、register/index.vue、stores/user.ts、utils/request.ts、utils/auth.ts 调用 authApi，并按 401/403 返回进行跳转或提示。",
        "api": "POST /api/auth/login；POST /api/auth/register；POST /api/auth/logout；GET /api/auth/me；POST /api/auth/refresh。",
        "summary": "该模块没有独立权限表或菜单表，权限模型来自 sys_user.role、is_approver 字段和前端静态路由 meta 配置，符合当前源码实现。"
    },
    {
        "title": "用户管理模块",
        "functions": "管理员维护员工信息，支持分页筛选、新增、修改、逻辑删除、批量删除、启停、重置密码、设置审批人身份；登录用户可维护个人资料。",
        "flow": "管理员进入用户管理页，前端按查询条件调用分页接口，后端使用 MyBatis-Plus LambdaQueryWrapper 组合用户名、姓名、手机号、部门、状态等条件。新增和重置密码时统一使用 BCrypt 加密，删除使用逻辑删除。",
        "classes": "UserController.page/create/update/delete/batchDelete/updateStatus/resetPassword/updateApprover/profile/updateProfile；UserServiceImpl.pageUsers/createUser/updateUser/resetPassword/updateStatus/updateApprover/listApprovers/updateProfile；UserMapper；User。",
        "tables": "sys_user：id、username、password、real_name、gender、phone、email、avatar、role、status、department_id、position、hire_date、is_approver、deleted。",
        "frontend": "frontend/src/views/user/index.vue、profile/index.vue、api/users.ts、components/PaginationTable.vue、components/OaUpload.vue。",
        "api": "GET /api/users/page；POST /api/users；PUT /api/users/{id}；DELETE /api/users/{id}；PUT /api/users/{id}/status；PUT /api/users/{id}/password/reset；PUT /api/users/{id}/approver；GET /api/users/profile。",
        "summary": "用户管理是审批、日程、通知等业务的人员基础，系统通过角色字段和审批人标记实现轻量权限划分。"
    },
    {
        "title": "部门组织管理模块",
        "functions": "维护组织结构，提供部门树、分页查询、部门新增修改删除、负责人和默认审批人设置。",
        "flow": "前端加载部门树并用于用户、审批、筛选等页面。管理员调整部门后，DepartmentServiceImpl 清理部门树缓存。提交审批时，审批服务通过部门默认审批人或负责人解析审批人。",
        "classes": "DepartmentController.tree/page/detail/create/update/delete/updateLeader/updateApprover；DepartmentServiceImpl.tree/createDepartment/updateDepartment/updateLeader/updateApprover/resolveApprover/clearTreeCache；DepartmentMapper；DepartmentTreeVO。",
        "tables": "sys_department：id、name、parent_id、leader_id、approver_id、sort_order、status、deleted；sys_user.department_id 关联所属部门。",
        "frontend": "frontend/src/views/department/index.vue、api/departments.ts，用户管理页也读取部门选项。",
        "api": "GET /api/departments/tree；GET /api/departments/page；POST /api/departments；PUT /api/departments/{id}；PUT /api/departments/{id}/leader；PUT /api/departments/{id}/approver。",
        "summary": "部门模块既承担组织展示，也承担审批人自动匹配的业务支撑。"
    },
    {
        "title": "审批流程管理模块",
        "functions": "支持请假、报销、加班、出差四类审批，包含草稿、新建、修改、提交、撤回、通过、驳回、审批记录、AI 摘要和风险提示。",
        "flow": "申请人创建草稿后提交，系统生成审批编号并根据部门解析审批人，状态变为 PENDING，同时写入审批记录并发送通知消息。审批人或管理员在待办中审批，通过或驳回后更新状态、记录意见、通知申请人。",
        "classes": "ApprovalController.page/my/todo/done/detail/create/update/submit/withdraw/approve/reject/records/aiSummary/aiRisk；ApprovalServiceImpl.submit/audit/ensureVisible/ensureCanApprove/generateApprovalNo；ApprovalMapper；ApprovalRecordMapper；ApprovalStatusEnum；ApprovalTypeEnum。",
        "tables": "oa_approval：approval_no、title、type、status、applicant_id、department_id、approver_id、reason、start_time、end_time、amount、destination、form_data、ai_summary、ai_risk_level；oa_approval_record：approval_id、operator_id、action、comment。",
        "frontend": "frontend/src/views/approval/list.vue、create.vue、todo.vue、detail.vue、api/approvals.ts。",
        "api": "GET /api/approvals/page；GET /api/approvals/my；GET /api/approvals/todo；POST /api/approvals；POST /api/approvals/{id}/submit；POST /api/approvals/{id}/approve；POST /api/approvals/{id}/reject；GET /api/approvals/{id}/records。",
        "summary": "审批为自研单级流程，源码中没有 Activiti、Flowable 或 BPMN 引擎。"
    },
    {
        "title": "日程会议管理模块",
        "functions": "支持个人日程、会议日程、参与人管理、接受/拒绝、今日/本周/日历视图、定时提醒和自然语言 AI 解析。",
        "flow": "用户创建日程时填写起止时间、地点、提醒分钟数和参与人。会议参与人记录保存在 oa_schedule_participant。定时任务每分钟扫描待提醒日程，发送 RabbitMQ 通知，消费者再生成站内信、WebSocket 推送并可发送邮件。",
        "classes": "ScheduleController.page/calendar/today/week/create/update/delete/addParticipants/accept/reject/aiParse；ScheduleServiceImpl.scanAndSendReminders/saveParticipants/updateParticipantStatus；ScheduleReminderTask；ScheduleMapper；ScheduleParticipantMapper。",
        "tables": "oa_schedule：title、content、type、creator_id、start_time、end_time、location、reminder_minutes、reminder_time、reminder_status、status、ai_origin_text；oa_schedule_participant：schedule_id、user_id、status。",
        "frontend": "frontend/src/views/schedule/list.vue、calendar.vue、api/schedules.ts。",
        "api": "GET /api/schedules/page；GET /api/schedules/calendar；POST /api/schedules；PUT /api/schedules/{id}；POST /api/schedules/{id}/participants；POST /api/schedules/{id}/accept；POST /api/schedules/ai-parse。",
        "summary": "日程模块把定时任务、消息队列、站内信、WebSocket 和邮件提醒串联成完整的办公提醒链路。"
    },
    {
        "title": "新闻公告管理模块",
        "functions": "支持新闻公告的新增、编辑、发布、下架、置顶、浏览、评论、点赞、收藏、我的收藏、AI 生成和 AI 润色。",
        "flow": "管理员在新闻管理页维护草稿，发布后状态变为 PUBLISHED 并通过通知链路提醒用户。普通员工在新闻列表中只查看已发布新闻，可进入详情页评论、点赞或收藏。新闻详情使用 Redis 缓存，修改、下架、删除后清理缓存。",
        "classes": "NewsController.page/detail/create/update/delete/publish/offline/top/view/comment/like/favorite/aiGenerate/aiPolish；NewsServiceImpl.pageNews/detail/publish/offline/comment/like/favorite/clearNewsCache；NewsMapper、NewsCommentMapper、NewsLikeMapper、NewsFavoriteMapper。",
        "tables": "oa_news、oa_news_comment、oa_news_like、oa_news_favorite，关键字段包括 status、is_top、view_count、like_count、favorite_count、comment_count、publisher_id。",
        "frontend": "frontend/src/views/news/list.vue、detail.vue、manage.vue、api/news.ts。",
        "api": "GET /api/news/page；GET /api/news/{id}；POST /api/news；POST /api/news/{id}/publish；POST /api/news/{id}/comments；POST /api/news/{id}/like；POST /api/news/{id}/favorite；POST /api/news/ai-generate。",
        "summary": "新闻模块同时覆盖公告发布、内容互动和 AI 辅助写作，是 OA 信息发布能力的核心。"
    },
    {
        "title": "字典与基础数据管理模块",
        "functions": "维护审批类型、审批状态、日程类型、新闻状态、通知类型等基础枚举数据，支持按类型读取和缓存刷新。",
        "flow": "管理员维护字典类型和字典项，前端表单通过 typeCode 获取可选项。DictDataServiceImpl 优先读取 Redis 缓存，字典新增、修改、删除后清理或刷新缓存。",
        "classes": "DictTypeController、DictDataController、DictTypeServiceImpl、DictDataServiceImpl、DictTypeMapper、DictDataMapper、CacheConstants。",
        "tables": "sys_dict_type：type_code、type_name、remark、status；sys_dict_data：type_code、dict_label、dict_value、sort_order、status。",
        "frontend": "frontend/src/views/dict/index.vue、components/DictSelect.vue、stores/dict.ts、api/dict.ts。",
        "api": "GET /api/dict-types/page；POST /api/dict-types；GET /api/dict-data/page；GET /api/dict-data/type/{typeCode}；POST /api/dict-data/cache/refresh。",
        "summary": "字典模块降低了前后端硬编码比例，并为审批、日程、新闻等模块提供统一选项来源。"
    },
    {
        "title": "文件上传与附件管理模块",
        "functions": "实现头像、审批附件、日程附件、新闻封面等文件上传、下载、列表查询、业务关联和逻辑删除。",
        "flow": "前端通过 OaUpload 组件上传 MultipartFile，后端按年月目录保存到本地磁盘，使用 UUID 重命名文件，并把原始文件名、存储路径、访问 URL、业务类型、业务 ID、上传人写入 sys_file。",
        "classes": "OaFileController.upload/detail/page/business/download/delete；OaFileServiceImpl.upload/download/deleteFile/ensureUploaderOrAdmin；FileProperties；WebMvcConfig。",
        "tables": "sys_file：original_name、file_name、file_path、file_url、file_size、file_type、extension、business_type、business_id、uploader_id。",
        "frontend": "frontend/src/components/OaUpload.vue、api/files.ts，个人中心、新闻管理、审批表单等页面复用。",
        "api": "POST /api/files/upload；GET /api/files/{id}；GET /api/files/business；GET /api/files/download/{id}；DELETE /api/files/{id}。",
        "summary": "源码使用本地磁盘存储和 /files 静态访问，没有实现 OSS 或云对象存储。"
    },
    {
        "title": "通知消息与实时推送模块",
        "functions": "提供站内信分页、未读数、标记已读、批量已读、删除、管理员系统通知，并串联 RabbitMQ、WebSocket 和邮件提醒。",
        "flow": "业务服务构造 NotificationMessage 后调用 NotificationProducer 投递到 oa.notification.exchange。NotificationConsumer 消费队列消息，写入 sys_notification，并调用 NotificationWebSocketHandler 给在线用户推送；日程和系统通知还会调用 MailService 发送邮件提醒。",
        "classes": "NotificationController；NotificationServiceImpl；NotificationProducer；NotificationConsumer；RabbitMqConfig；NotificationWebSocketHandler；NotificationHandshakeInterceptor；WebSocketConfig；MailServiceImpl。",
        "tables": "sys_notification：receiver_id、sender_id、title、content、type、business_type、business_id、read_status、read_time、pushed。",
        "frontend": "frontend/src/views/notification/index.vue、components/NotificationBell.vue、stores/notification.ts、utils/ws.ts。",
        "api": "GET /api/notifications/page；GET /api/notifications/unread-count；PUT /api/notifications/{id}/read；PUT /api/notifications/read-batch；POST /api/notifications/system；WebSocket /ws/notification?token=xxx。",
        "summary": "通知模块是系统异步化和实时化的关键支撑，体现了消息队列与 WebSocket 在 OA 场景中的协作。"
    },
    {
        "title": "AI 助手与调用日志模块",
        "functions": "提供审批摘要、审批风险分析、新闻生成、新闻润色、日程解析、智能问答，并记录 AI 调用日志。",
        "flow": "前端 AI 页面或业务详情页调用 AI 接口，AiController 根据功能调用 AiService。系统支持 MockAiServiceImpl 和 TongyiAiServiceImpl，可通过 oa.ai.provider 切换；每次调用写入 sys_ai_log 记录功能类型、请求、响应、耗时和成功状态。",
        "classes": "AiController；AiService；MockAiServiceImpl；TongyiAiServiceImpl；AiLogServiceImpl；AiProperties；AiLogMapper；AiFunctionTypeEnum。",
        "tables": "sys_ai_log：user_id、function_type、provider、model_name、prompt、request_content、response_content、success、error_message、cost_time_ms。",
        "frontend": "frontend/src/views/ai/index.vue、components/AiDialog.vue、api/ai.ts。",
        "api": "POST /api/ai/approval-summary；POST /api/ai/approval-risk；POST /api/ai/news-generate；POST /api/ai/news-polish；POST /api/ai/schedule-parse；POST /api/ai/qa；GET /api/ai/logs/page。",
        "summary": "系统没有独立操作日志表，当前可确认的结构化日志是 AI 调用日志 sys_ai_log；其他系统异常通过后端日志输出。"
    },
]

COURSE_TOC = [
    (1, "绪论"),
    (1, "1 需求分析与设计"),
    (2, "1.1 系统需求分析"),
    (2, "1.2 系统设计方案"),
    (2, "1.3 功能分析"),
    (2, "1.4 数据流程分析"),
    (2, "1.5 业务流程分析"),
    (2, "1.6 系统实现技术选型"),
    (1, "2 系统设计"),
    (2, "2.1 系统概要设计"),
    (2, "2.2 系统结构设计"),
    (2, "2.3 数据库设计"),
    (3, "2.3.1 数据库表设计"),
    (1, "3 系统实现"),
    *[(2, f"3.{i} {m['title']}") for i, m in enumerate(MODULES, 1)],
    (1, "4 系统测试"),
    (2, "4.1 测试定义"),
    (2, "4.2 测试目的"),
    (2, "4.3 测试方法"),
    (2, "4.4 测试用例"),
    (2, "4.5 测试分析"),
    (1, "5 总结"),
]

INITIATION_TOC = [
    (1, "1 Project Proposal 项目提出"),
    (2, "1.1 Project Brief 项目简介"),
    (2, "1.2 Project Goal 项目目标"),
    (2, "1.3 System Scope 系统边界"),
    (2, "1.4 Estimated Effort 工作量估计"),
    (1, "2 Team building and Schedule 开发团队组成和计划时间"),
    (2, "2.1 Project Team 开发团队"),
    (2, "2.2 Project Plan 计划时间"),
    (1, "3 Evaluating and Mitigating 风险评估和规避"),
    (2, "3.1 Technical Risks 技术风险"),
    (2, "3.2 Management Risks 管理风险"),
]

REQUIREMENTS_TOC = [
    (1, "1 Introduction 简介"),
    (2, "1.1 Purpose 文档目的"),
    (2, "1.2 Scope 本文档适用范围"),
    (1, "2 General description 总体概述"),
    (2, "2.1 Soft perspective 软件概述"),
    (3, "2.1.1 About the Project 项目介绍"),
    (2, "2.2 Soft function 软件功能"),
    (2, "2.3 User Roles and Permissions 用户角色与权限"),
    (1, "3 Functional Requirements 功能需求"),
    *[(2, f"3.{i} {m['title']}") for i, m in enumerate(MODULES, 1)],
    (1, "4 Performance Requirements 非功能需求"),
    (2, "4.1 UI Requirements 界面要求"),
    (2, "4.2 Development Environment 开发环境"),
    (2, "4.3 Development Rules 开发规范"),
    (1, "5 Demand Classification 需求分级"),
]

DESIGN_TOC = [
    (1, "1 Introduction  简介"),
    (2, "1.1 Purpose  文档目的"),
    (2, "1.2 Scope  本文档适用范围"),
    (2, "1.3 Name 软件名称"),
    (2, "1.4 Applications 软件应用领域"),
    (1, "2 High Level Design 概要设计"),
    (2, "2.1 系统功能设计"),
    (2, "2.2 系统架构设计"),
    (3, "2.2.1 前后端分离架构"),
    (3, "2.2.2 认证与权限架构"),
    (3, "2.2.3 通知、RabbitMQ 与 WebSocket 链路"),
    (2, "2.3 Database 数据库设计"),
    (3, "2.3.1 表关系"),
    (3, "2.3.2 数据表设计"),
    (1, "3 Low Level Model Design 模块详细设计"),
    *[(2, f"3.{i} {m['title']}") for i, m in enumerate(MODULES, 1)],
]

SCREENSHOTS = {
    "login": ROOT / "frontend" / ".qa-full" / "01-login-desktop.png",
    "dashboard": ROOT / "frontend" / ".qa-full" / "02-dashboard-desktop.png",
    "users": ROOT / "frontend" / ".qa-screens" / "users.png",
    "departments": ROOT / "frontend" / ".qa-screens" / "departments.png",
    "approvals": ROOT / "frontend" / ".qa-screens" / "approvals.png",
    "approval_create": ROOT / "frontend" / ".qa-screens" / "approvals_create.png",
    "approval_todo": ROOT / "frontend" / ".qa-screens" / "approvals_todo.png",
    "schedules": ROOT / "frontend" / ".qa-screens" / "schedules.png",
    "calendar": ROOT / "frontend" / ".qa-screens" / "schedules_calendar.png",
    "news": ROOT / "frontend" / ".qa-screens" / "news.png",
    "news_manage": ROOT / "frontend" / ".qa-screens" / "news_manage.png",
    "dicts": ROOT / "frontend" / ".qa-screens" / "dicts.png",
    "notifications": ROOT / "frontend" / ".qa-screens" / "notifications.png",
    "ai": ROOT / "frontend" / ".qa-screens" / "ai.png",
    "profile": ROOT / "frontend" / ".qa-screens" / "profile.png",
}


def ensure_dirs() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_HEI if bold else FONT_SONG
    return ImageFont.truetype(path, size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, ft: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    line = ""
    for ch in text:
        trial = line + ch
        if draw.textbbox((0, 0), trial, font=ft)[2] <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = ch
    if line:
        lines.append(line)
    return lines


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str, outline: str = "#334155", ft=None) -> None:
    ft = ft or font(28, True)
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    lines = wrap_text(draw, text, ft, x2 - x1 - 30)
    line_h = ft.size + 8
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=ft)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) // 2, y), line, fill="#0f172a", font=ft)
        y += line_h


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#2563eb", width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    import math
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (end[0] - size * math.cos(ang - 0.45), end[1] - size * math.sin(ang - 0.45))
    p2 = (end[0] - size * math.cos(ang + 0.45), end[1] - size * math.sin(ang + 0.45))
    draw.polygon([end, p1, p2], fill=color)


def new_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 38), title, fill="#111827", font=font(42, True))
    draw.line((70, 95, 1530, 95), fill="#cbd5e1", width=3)
    return img, draw


def make_diagrams() -> dict[str, Path]:
    diagrams: dict[str, Path] = {}

    img, draw = new_canvas("OA_SYSTEM 系统总体架构图")
    draw_box(draw, (90, 160, 390, 280), "Vue3 管理端\nElement Plus\nPinia Router Axios", "#dbeafe")
    draw_box(draw, (590, 140, 1010, 310), "Spring Boot 3 后端\nController Service Mapper\n统一返回与异常处理", "#dcfce7")
    draw_box(draw, (1160, 130, 1510, 250), "AI 接入\nMock / 通义千问\nsys_ai_log", "#fef3c7")
    arrow(draw, (390, 220), (590, 220))
    draw.text((440, 185), "REST API / JWT", fill="#1d4ed8", font=font(24, True))
    arrow(draw, (390, 260), (590, 280), "#7c3aed")
    draw.text((430, 285), "WebSocket 通知", fill="#6d28d9", font=font(22, True))
    arrow(draw, (1010, 220), (1160, 190), "#f59e0b")
    middles = [
        ("MySQL 8\noa_management", 120, 560, "#fee2e2"),
        ("Redis\n字典/部门/新闻/未读缓存", 440, 560, "#e0f2fe"),
        ("RabbitMQ\nTopicExchange 通知队列", 810, 560, "#f3e8ff"),
        ("本地文件目录\nbackend/upload", 1190, 560, "#ecfccb"),
    ]
    for text, x, y, fill in middles:
        draw_box(draw, (x, y, x + 290, y + 145), text, fill, ft=font(24, True))
        arrow(draw, (805, 310), (x + 145, y), "#64748b", 3)
    img.save(ASSETS / "system_architecture.png")
    diagrams["系统总体架构图"] = ASSETS / "system_architecture.png"

    img, draw = new_canvas("OA_SYSTEM 系统功能模块图")
    draw_box(draw, (640, 370, 960, 500), "企业 OA 管理系统\nOA_SYSTEM", "#e2e8f0", ft=font(30, True))
    items = [
        ("登录认证\nJWT 权限", 120, 160),
        ("用户管理\n员工资料", 440, 150),
        ("部门管理\n组织树", 760, 150),
        ("审批流程\n请假报销加班出差", 1080, 160),
        ("日程会议\n提醒与参与人", 1260, 390),
        ("新闻公告\n评论点赞收藏", 1080, 650),
        ("字典配置\n基础数据", 760, 670),
        ("文件附件\n上传下载", 440, 670),
        ("通知与 AI\nWebSocket / 日志", 120, 650),
    ]
    for text, x, y in items:
        draw_box(draw, (x, y, x + 230, y + 120), text, "#f8fafc", ft=font(22, True))
        arrow(draw, (800, 435), (x + 115, y + 60), "#475569", 3)
    img.save(ASSETS / "function_modules.png")
    diagrams["系统功能模块图"] = ASSETS / "function_modules.png"

    img, draw = new_canvas("OA_SYSTEM 数据库 ER 概要图")
    pos = {
        "sys_user": (90, 150), "sys_department": (430, 150), "oa_approval": (760, 150), "oa_approval_record": (1110, 150),
        "oa_schedule": (90, 390), "oa_schedule_participant": (430, 390), "oa_news": (760, 390), "sys_file": (1110, 390),
        "sys_notification": (90, 640), "sys_dict_type": (430, 640), "sys_dict_data": (760, 640), "sys_ai_log": (1110, 640),
    }
    labels = {
        "sys_user": "sys_user\n用户/角色/审批人",
        "sys_department": "sys_department\n部门/负责人",
        "oa_approval": "oa_approval\n审批主表",
        "oa_approval_record": "oa_approval_record\n审批记录",
        "oa_schedule": "oa_schedule\n日程主表",
        "oa_schedule_participant": "oa_schedule_participant\n参与人",
        "oa_news": "oa_news\n新闻公告",
        "sys_file": "sys_file\n文件元数据",
        "sys_notification": "sys_notification\n站内信",
        "sys_dict_type": "sys_dict_type\n字典类型",
        "sys_dict_data": "sys_dict_data\n字典数据",
        "sys_ai_log": "sys_ai_log\nAI 日志",
    }
    for key, (x, y) in pos.items():
        draw_box(draw, (x, y, x + 250, y + 110), labels[key], "#f8fafc", ft=font(21, True))
    relations = [
        ("sys_department", "sys_user"), ("sys_user", "oa_approval"), ("oa_approval", "oa_approval_record"),
        ("sys_user", "oa_schedule"), ("oa_schedule", "oa_schedule_participant"), ("sys_user", "oa_schedule_participant"),
        ("sys_user", "oa_news"), ("oa_news", "sys_file"), ("sys_user", "sys_notification"),
        ("sys_dict_type", "sys_dict_data"), ("sys_user", "sys_ai_log"),
    ]
    for a, b in relations:
        ax, ay = pos[a]; bx, by = pos[b]
        arrow(draw, (ax + 125, ay + 110), (bx + 125, by), "#64748b", 3)
    draw.text((90, 820), "说明：schema.sql 未定义数据库外键约束，图中关系依据业务字段和 Service 逻辑整理。", fill="#334155", font=font(22, False))
    img.save(ASSETS / "database_er.png")
    diagrams["数据库 ER 概要图"] = ASSETS / "database_er.png"

    def flow_diagram(name: str, title: str, steps: list[str], color: str) -> None:
        img, draw = new_canvas(title)
        x0, y, w, h = 90, 330, 220, 120
        gap = 35
        for i, step in enumerate(steps):
            x = x0 + i * (w + gap)
            draw_box(draw, (x, y, x + w, y + h), step, "#f8fafc", color, ft=font(21, True))
            if i < len(steps) - 1:
                arrow(draw, (x + w, y + h // 2), (x + w + gap, y + h // 2), color, 4)
        img.save(ASSETS / name)

    flow_diagram(
        "login_flow.png",
        "登录认证流程图",
        ["登录页输入账号密码", "AuthController\n接收请求", "UserService\n校验用户与密码", "JwtTokenUtil\n生成 JWT", "前端保存 Token\n后续请求携带", "JwtFilter\n解析并鉴权"],
        "#2563eb",
    )
    diagrams["登录认证流程图"] = ASSETS / "login_flow.png"

    flow_diagram(
        "approval_flow.png",
        "审批业务流程图",
        ["创建审批草稿", "提交审批\n生成编号", "匹配部门审批人", "写入审批记录\n发送通知", "审批人处理", "更新状态\n通知申请人"],
        "#16a34a",
    )
    diagrams["审批业务流程图"] = ASSETS / "approval_flow.png"

    flow_diagram(
        "notification_flow.png",
        "通知消息流转图",
        ["业务事件", "NotificationProducer", "RabbitMQ\nTopicExchange", "NotificationConsumer", "站内信落库", "WebSocket 推送\n邮件提醒"],
        "#7c3aed",
    )
    diagrams["通知消息流转图"] = ASSETS / "notification_flow.png"

    img, draw = new_canvas("OA_SYSTEM 部署结构图")
    draw_box(draw, (90, 170, 360, 300), "浏览器用户\n管理端页面", "#dbeafe")
    draw_box(draw, (570, 170, 930, 310), "Vue3 前端\n开发端口 5173\n或 dist 静态资源", "#e0f2fe")
    draw_box(draw, (1120, 170, 1510, 310), "Spring Boot JAR\n端口 8081\n/api 与 /ws", "#dcfce7")
    arrow(draw, (360, 235), (570, 235))
    arrow(draw, (930, 235), (1120, 235))
    services = [
        ("MySQL 8\n3306\noa_management", 170, 560, "#fee2e2"),
        ("Redis 7\n6379\n缓存与黑名单", 500, 560, "#fef3c7"),
        ("RabbitMQ\n6672 / 15672\n通知队列", 830, 560, "#f3e8ff"),
        ("upload 目录\n文件附件", 1160, 560, "#ecfccb"),
    ]
    for text, x, y, fill in services:
        draw_box(draw, (x, y, x + 270, y + 135), text, fill, ft=font(23, True))
        arrow(draw, (1315, 310), (x + 135, y), "#64748b", 3)
    img.save(ASSETS / "deployment_structure.png")
    diagrams["系统部署结构图"] = ASSETS / "deployment_structure.png"

    return diagrams


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.7)
    sec.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing = 1.25

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        st = doc.styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.paragraph_format.first_line_indent = Pt(0)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc: Document, headers: list[str], rows: list[Iterable[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = header_cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, 10, True)
        shade_cell(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, 9.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(r, 10, False)


def add_image(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if not path.exists():
        add_paragraph(doc, f"图片文件未找到：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_toc(doc: Document, entries: list[tuple[int, str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("目  录")
    set_run_font(r, 16, True)
    for level, text in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt((level - 1) * 18)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        set_run_font(r, 10.5, level == 1)


def add_cover(doc: Document, doc_title: str, english_title: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(english_title)
    set_run_font(r, 18, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(doc_title)
    set_run_font(r, 22, True)

    for _ in range(3):
        doc.add_paragraph()

    info = [
        ("项目名称", PROJECT_CN, "系统标识", PROJECT_EN),
        ("班级", CLASS_NAME, "指导老师", TEACHER),
        ("组长", "何成翔", "日期", ""),
    ]
    add_table(doc, ["项目项", "内容", "项目项", "内容"], info, [3, 5, 3, 5])

    add_table(
        doc,
        ["姓名", "小组职责", "学号"],
        TEAM,
        [4, 5, 5],
    )
    doc.add_page_break()


def add_common_project_intro(doc: Document) -> None:
    add_paragraph(
        doc,
        "OA_SYSTEM 是一个面向企业日常办公场景的前后端分离管理系统。系统围绕用户、部门、审批、日程、新闻公告、字典、文件、站内信通知和 AI 助手等业务展开，后端提供统一 RESTful JSON API，前端提供基于 Vue3 和 Element Plus 的管理端界面。",
    )
    add_paragraph(
        doc,
        "经源码核查，本项目后端为单体 Spring Boot 模块化应用，并非多个独立后端服务组成的 Spring Cloud 集群。文档中的微服务架构表述仅用于说明课程主题下的分层、模块化、缓存、消息队列、实时推送和容器化依赖协作，不写入源码中不存在的网关、注册中心和配置中心。"
    )


def add_module_impl(doc: Document, module: dict[str, str]) -> None:
    doc.add_heading(module["title"], 2)
    add_table(
        doc,
        ["说明项", "内容"],
        [
            ("模块功能说明", module["functions"]),
            ("业务流程说明", module["flow"]),
            ("核心类/接口/方法", module["classes"]),
            ("数据库表或关键字段", module["tables"]),
            ("前后端交互逻辑", module["frontend"]),
            ("接口调用效果", module["api"]),
            ("模块实现小结", module["summary"]),
        ],
        [4, 12],
    )


def build_course_report(diagrams: dict[str, Path]) -> Path:
    doc = Document()
    style_doc(doc)
    add_cover(doc, "SpringBoot 微服务架构课程设计报告\n企业 OA 管理系统", "SpringBoot Microservice Architecture Course Design Report")
    add_toc(doc, COURSE_TOC)
    doc.add_page_break()

    doc.add_heading("绪论", 1)
    add_common_project_intro(doc)
    add_paragraph(
        doc,
        "随着企业办公流程数字化程度不断提高，传统依赖人工流转的请假、报销、会议通知和公告发布方式容易出现信息滞后、数据分散、审批记录难追溯等问题。本课程设计选择企业 OA 管理系统作为实现对象，重点完成用户认证、组织管理、审批、日程、新闻公告、通知推送、文件附件和 AI 辅助办公等功能，并通过 Redis、RabbitMQ、WebSocket 等技术提升系统的响应速度与实时协同能力。"
    )
    add_image(doc, diagrams["系统总体架构图"], "图 0.1 OA_SYSTEM 系统总体架构图")

    doc.add_heading("1 需求分析与设计", 1)
    doc.add_heading("1.1 系统需求分析", 2)
    add_paragraph(doc, "系统目标用户包括管理员、普通员工和具备审批身份的员工。管理员负责维护用户、部门、字典、新闻、通知等基础数据；普通员工负责发起审批、管理个人日程、查看公告、处理个人资料和站内信；审批人员在普通员工能力基础上可处理分配给自己的待办审批。")
    add_table(doc, ["用户角色", "主要需求", "权限边界"], [
        ("管理员 ADMIN", "用户、部门、字典、新闻、通知、AI 日志等管理；可查看全部审批和系统数据。", "通过 hasRole('ADMIN') 控制管理接口。"),
        ("员工 EMPLOYEE", "登录系统、维护个人资料、发起审批、管理日程、查看新闻、接收通知、使用部分 AI 功能。", "只能访问本人或参与范围内的数据。"),
        ("审批员工", "处理待本人审批的审批单，并可使用审批相关辅助功能。", "由 sys_user.is_approver 字段和 SecurityPermission 判断。"),
    ], [3.5, 7, 5])

    doc.add_heading("1.2 系统设计方案", 2)
    add_paragraph(doc, "系统采用前后端分离方案。前端负责界面展示、路由控制、状态管理和用户交互；后端负责认证授权、业务规则、数据访问、缓存、消息投递、文件保存和 AI 调用；MySQL 存储业务数据，Redis 缓存高频数据和 Token 黑名单，RabbitMQ 处理通知异步事件，WebSocket 负责在线用户实时推送。")
    add_image(doc, diagrams["系统部署结构图"], "图 1.1 OA_SYSTEM 系统部署结构图")

    doc.add_heading("1.3 功能分析", 2)
    add_image(doc, diagrams["系统功能模块图"], "图 1.2 OA_SYSTEM 系统功能模块图")
    add_table(doc, ["模块", "功能范围"], [(m["title"], m["functions"]) for m in MODULES], [4.5, 11.5])

    doc.add_heading("1.4 数据流程分析", 2)
    add_paragraph(doc, "系统数据流以登录 Token 为访问入口。前端请求首先经过 Axios 拦截器附加 JWT，后端过滤器完成认证后进入 Controller，Controller 将参数交给 Service 执行业务校验，Mapper 通过 MyBatis-Plus 访问数据库。对于高频读取数据，Service 优先访问 Redis；对于通知类事件，Service 不直接完成全部后续处理，而是投递 RabbitMQ 消息，由消费者落库并推送。")
    add_image(doc, diagrams["通知消息流转图"], "图 1.3 通知消息数据流转图")

    doc.add_heading("1.5 业务流程分析", 2)
    add_paragraph(doc, "本系统最具代表性的业务流程为登录认证、审批流转和通知推送。登录认证保障后续接口访问的身份可信；审批流转体现部门、用户和通知模块之间的协同；通知推送体现异步消息、站内信和 WebSocket 的组合使用。")
    add_image(doc, diagrams["登录认证流程图"], "图 1.4 登录认证流程图")
    add_image(doc, diagrams["审批业务流程图"], "图 1.5 审批业务流程图")

    doc.add_heading("1.6 系统实现技术选型", 2)
    add_table(doc, ["类别", "实际采用技术"], TECH_STACK, [5, 11])
    for title, text in [
        ("1.6.1 Spring Boot 与 Spring Security 技术", "后端基于 Spring Boot 3.2.5 构建，使用嵌入式 Web 容器和注解式开发方式。Spring Security 负责认证授权，JWT 作为无状态登录凭证，BCrypt 用于密码加密。"),
        ("1.6.2 Vue3 与 Element Plus 技术", "前端使用 Vue3 组合式 API、Vite 构建、Element Plus 组件库、Pinia 状态管理和 Vue Router 路由。"),
        ("1.6.3 MySQL 与 MyBatis-Plus 技术", "MySQL 保存业务数据，MyBatis-Plus 提供通用 CRUD、分页、逻辑删除和字段自动填充能力。"),
        ("1.6.4 Redis、RabbitMQ、WebSocket 与 AI 接入技术", "Redis 用于字典、部门树、新闻和未读数缓存；RabbitMQ 实现异步通知；WebSocket 推送在线消息；AI 模块支持 mock 与通义千问提供商切换。"),
    ]:
        doc.add_heading(title, 3)
        add_paragraph(doc, text)

    doc.add_heading("2 系统设计", 1)
    doc.add_heading("2.1 系统概要设计", 2)
    add_paragraph(doc, "系统按 Controller、Service、Mapper、Entity/DTO/VO 分层实现。Controller 负责接口入口和参数校验，Service 负责业务规则与事务，Mapper 负责数据访问，公共层提供统一返回、分页、异常处理、缓存常量、安全工具和通用实体字段。")
    doc.add_heading("2.2 系统结构设计", 2)
    add_table(doc, ["源码目录", "作用"], [
        ("backend/src/main/java/com/example/oa/common", "通用实体、统一返回、分页、异常、枚举、常量和安全工具。"),
        ("backend/src/main/java/com/example/oa/config", "MyBatis-Plus、Redis、RabbitMQ、WebSocket、文件访问、OpenAPI、异步调度配置。"),
        ("backend/src/main/java/com/example/oa/security", "Spring Security、JWT、登录用户、权限判断、认证和授权异常处理。"),
        ("backend/src/main/java/com/example/oa/module", "按用户、部门、审批、日程、新闻、字典、文件、通知、AI 划分业务模块。"),
        ("frontend/src/views", "登录、仪表盘、用户、部门、审批、日程、新闻、字典、通知、AI、个人中心页面。"),
        ("frontend/src/api", "前端接口封装，与后端 /api 路径对应。"),
    ], [6, 10])
    doc.add_heading("2.3 数据库设计", 2)
    add_image(doc, diagrams["数据库 ER 概要图"], "图 2.1 OA_SYSTEM 数据库 ER 概要图")
    doc.add_heading("2.3.1 数据库表设计", 3)
    add_table(doc, ["表名", "中文说明", "主要字段或作用"], TABLES, [4.5, 4, 8])

    doc.add_heading("3 系统实现", 1)
    add_paragraph(doc, "本章依据源码实际模块展开。对于课程要求中列出的网关、注册中心、配置中心、独立菜单权限表等内容，源码中未实现，因此不作为已完成功能编写。")
    for idx, module in enumerate(MODULES, 1):
        add_module_impl(doc, module)
        if idx == 1:
            add_image(doc, SCREENSHOTS["login"], "图 3.1 登录页面真实截图")
        elif idx == 2:
            add_image(doc, SCREENSHOTS["users"], "图 3.2 用户管理页面真实截图")
        elif idx == 3:
            add_image(doc, SCREENSHOTS["departments"], "图 3.3 部门管理页面真实截图")
        elif idx == 4:
            add_image(doc, SCREENSHOTS["approvals"], "图 3.4 审批列表页面真实截图")
            add_image(doc, SCREENSHOTS["approval_todo"], "图 3.5 待我审批页面真实截图")
        elif idx == 5:
            add_image(doc, SCREENSHOTS["schedules"], "图 3.6 日程列表页面真实截图")
            add_image(doc, SCREENSHOTS["calendar"], "图 3.7 日程日历页面真实截图")
        elif idx == 6:
            add_image(doc, SCREENSHOTS["news_manage"], "图 3.8 新闻管理页面真实截图")
        elif idx == 7:
            add_image(doc, SCREENSHOTS["dicts"], "图 3.9 字典管理页面真实截图")
        elif idx == 9:
            add_image(doc, SCREENSHOTS["notifications"], "图 3.10 通知中心页面真实截图")
        elif idx == 10:
            add_image(doc, SCREENSHOTS["ai"], "图 3.11 AI 助手页面真实截图")

    doc.add_heading("4 系统测试", 1)
    doc.add_heading("4.1 测试定义", 2)
    add_paragraph(doc, "系统测试是对已实现功能、接口、权限、页面交互和部署运行环境进行验证的过程，目的是确认系统是否满足课程设计目标和项目需求。")
    doc.add_heading("4.2 测试目的", 2)
    add_paragraph(doc, "测试重点包括登录认证是否可靠、管理员与员工权限是否隔离、审批状态流转是否正确、RabbitMQ 与 WebSocket 通知链路是否可用、前端页面是否能调用真实接口并展示数据。")
    doc.add_heading("4.3 测试方法", 2)
    add_bullets(doc, [
        "代码级检查：核对 Controller、Service、Mapper、实体类和 SQL 表结构。",
        "接口级检查：依据 docs/api-overview.md 和控制器路径验证 REST API。",
        "页面级检查：使用已保存的真实页面截图核对主要功能页面。",
        "构建检查：后端 Maven 打包和前端 npm run build。",
    ])
    doc.add_heading("4.4 测试用例", 2)
    add_table(doc, ["编号", "测试项", "测试步骤", "预期结果"], [
        ("TC-01", "登录认证", "使用管理员账号登录，访问 /api/auth/me。", "返回当前用户信息，后续请求携带 Bearer Token。"),
        ("TC-02", "权限控制", "员工访问管理员用户管理接口。", "返回 403 或前端不展示对应菜单。"),
        ("TC-03", "用户管理", "管理员新增、修改、禁用、重置员工密码。", "sys_user 数据变化正确，密码使用 BCrypt。"),
        ("TC-04", "审批流程", "员工创建审批、提交，审批人通过或驳回。", "审批状态与 oa_approval_record 记录正确。"),
        ("TC-05", "通知推送", "提交审批或发送系统通知。", "sys_notification 落库，在线用户可收到 WebSocket 推送。"),
        ("TC-06", "日程提醒", "创建带提醒时间的日程。", "定时任务扫描后投递通知消息。"),
        ("TC-07", "新闻公告", "管理员发布新闻，员工评论、点赞、收藏。", "新闻状态和计数字段正确更新。"),
        ("TC-08", "AI 助手", "调用日程解析或智能问答。", "返回结构化响应并写入 sys_ai_log。"),
    ], [2, 3, 7, 6])
    doc.add_heading("4.5 测试分析", 2)
    add_paragraph(doc, "从源码、页面截图和构建产物看，系统已覆盖课程设计所需的主要办公业务。需要人工复核的内容主要是实际运行环境中的数据库账号、RabbitMQ 端口与 application.yml 是否一致，以及通义千问 API Key 是否已配置。")

    doc.add_heading("5 总结", 1)
    add_paragraph(doc, "本课程设计完成了一个面向企业办公场景的 OA 管理系统。系统在业务上覆盖用户、部门、审批、日程、新闻、字典、文件、通知和 AI 助手；在技术上使用 Spring Boot、Spring Security、MyBatis-Plus、MySQL、Redis、RabbitMQ、WebSocket 和 Vue3 等技术完成前后端分离实现。")
    add_paragraph(doc, "项目的不足在于当前后端不是 Spring Cloud 多服务集群，未实现独立网关、注册中心、配置中心、菜单权限表和完整操作日志表。后续可在现有模块化基础上继续拆分认证、通知、文件、AI 等服务，并补充操作审计、死信队列、部署脚本和自动化测试。")

    path = OUT / "SpringBoot微服务架构课程设计报告_OA_SYSTEM.docx"
    doc.save(path)
    return path


def build_initiation_doc(diagrams: dict[str, Path]) -> Path:
    doc = Document()
    style_doc(doc)
    add_cover(doc, "企业 OA 管理系统项目立项说明书", "Project Start Report")
    add_toc(doc, INITIATION_TOC)
    doc.add_page_break()
    doc.add_heading("1 Project Proposal 项目提出", 1)
    doc.add_heading("1.1 Project Brief 项目简介", 2)
    add_common_project_intro(doc)
    doc.add_heading("1.2 Project Goal 项目目标", 2)
    add_bullets(doc, [
        "实现前后端分离的企业 OA 管理端，提升办公事项线上处理能力。",
        "完成用户认证、权限控制、用户管理、部门管理、审批、日程、新闻、字典、文件、通知和 AI 助手等模块。",
        "通过 Redis、RabbitMQ、WebSocket 和邮件提醒形成缓存、异步通知和实时推送能力。",
        "形成可提交课程设计的源码、数据库脚本、项目说明书和课程设计报告。"
    ])
    doc.add_heading("1.3 System Scope 系统边界", 2)
    add_image(doc, diagrams["系统功能模块图"], "图 1.1 OA_SYSTEM 功能结构图")
    add_paragraph(doc, "系统边界包括管理端 Web 页面、Spring Boot 后端 API、MySQL 数据库、Redis 缓存、RabbitMQ 消息队列、WebSocket 通知通道、本地文件存储和 AI 服务接入。不包含源码未实现的 Spring Cloud Gateway、Nacos/Eureka、配置中心、独立角色菜单权限表和多后端服务拆分。")
    doc.add_heading("1.4 Estimated Effort 工作量估计", 2)
    add_table(doc, ["模块", "子模块", "工作量估计（工时）", "说明"], [
        ("工程基础", "后端/前端工程、Docker Compose", "20", "建立 Spring Boot、Vue3、MySQL、Redis、RabbitMQ 基础环境。"),
        ("认证与用户", "登录、JWT、用户管理、个人中心", "28", "实现账号、角色、审批人身份和资料维护。"),
        ("组织与字典", "部门树、字典类型/数据", "20", "支撑审批人匹配和表单选项。"),
        ("审批业务", "草稿、提交、审批、记录、AI 辅助", "36", "核心办公流程模块。"),
        ("日程业务", "日程、会议、参与人、提醒、AI 解析", "32", "含定时任务和通知联动。"),
        ("新闻公告", "发布、下架、互动、缓存、AI 写作", "30", "覆盖公告和内容互动。"),
        ("通知文件", "站内信、WebSocket、邮件、上传下载", "32", "消息队列、实时推送和附件管理。"),
        ("前端页面", "布局、路由、表格、表单、截图验证", "40", "完成管理端主要页面和交互。"),
        ("文档与测试", "报告、说明书、构建检查", "24", "整理交付文档和验收材料。"),
    ], [3, 4, 3, 8])
    doc.add_heading("2 Team building and Schedule 开发团队组成和计划时间", 1)
    doc.add_heading("2.1 Project Team 开发团队", 2)
    add_table(doc, ["Team 团队成员", "Name 姓名", "Student ID 学号"], [(role, name, sid) for name, role, sid in TEAM], [5, 5, 5])
    doc.add_heading("2.2 Project Plan 计划时间", 2)
    add_table(doc, ["阶段", "主要工作", "交付物"], [
        ("需求与设计", "确认课程要求、项目模块、数据库表和接口契约。", "需求摘要、接口概览、数据库设计。"),
        ("基础开发", "搭建后端、前端和 Docker Compose 依赖。", "可启动工程和基础配置。"),
        ("业务开发", "实现用户、部门、审批、日程、新闻、字典、文件、通知、AI 模块。", "完整源码和数据库脚本。"),
        ("联调测试", "进行前后端联调、页面截图、构建测试。", "测试截图、构建产物、问题修正。"),
        ("文档交付", "生成课程设计报告和三份说明书。", "DOCX 文档和质量检查结果。"),
    ], [3, 8, 6])
    doc.add_heading("3 Evaluating and Mitigating 风险评估和规避", 1)
    doc.add_heading("3.1 Technical Risks 技术风险", 2)
    add_table(doc, ["风险", "影响", "规避措施"], [
        ("RabbitMQ 或 Redis 端口配置不一致", "通知和缓存功能无法正常联调。", "以 application.yml 和 docker-compose.yml 为准检查端口，启动前统一配置。"),
        ("AI Key 未配置", "通义千问 provider 无法真实调用。", "保留 MockAiServiceImpl，真实环境再配置 DASHSCOPE_API_KEY。"),
        ("本地文件目录权限不足", "头像或附件无法保存。", "确认 backend/upload 目录可写，并在部署时配置持久化目录。"),
        ("权限边界遗漏", "员工可能访问管理数据。", "接口使用 @PreAuthorize 和业务层 ensureVisible/ensureMine 校验。"),
    ], [5, 6, 7])
    doc.add_heading("3.2 Management Risks 管理风险", 2)
    add_table(doc, ["风险", "影响", "规避措施"], [
        ("文档与源码不一致", "课程提交时被判定为空泛或编造。", "所有模块、表名、接口均从源码、SQL 和截图核对。"),
        ("小组分工交接不清", "重复开发或遗漏模块。", "按用户/部门、审批、日程、新闻、通知、前端、文档分块交付。"),
        ("测试环境不可复现", "验收时无法展示运行效果。", "保留 docker-compose、初始化 SQL、构建命令和真实截图。"),
    ], [5, 6, 7])
    path = OUT / "企业OA管理系统项目立项说明书—程序员.docx"
    doc.save(path)
    return path


def build_requirements_doc(diagrams: dict[str, Path]) -> Path:
    doc = Document()
    style_doc(doc)
    add_cover(doc, "企业 OA 管理系统软件需求说明书", "Soft Requirements Specification")
    add_toc(doc, REQUIREMENTS_TOC)
    doc.add_page_break()
    doc.add_heading("1 Introduction 简介", 1)
    doc.add_heading("1.1 Purpose 文档目的", 2)
    add_paragraph(doc, "本文档用于说明企业 OA 管理系统的功能需求、非功能需求、用户角色、接口边界和需求优先级，为开发、测试和课程验收提供统一依据。")
    doc.add_heading("1.2 Scope 本文档适用范围", 2)
    add_paragraph(doc, "本文档适用于 OA_SYSTEM 项目的需求说明、开发实现、功能验收和后续维护，不适用于源码中未实现的 Spring Cloud 网关、注册中心、配置中心或外部对象存储系统。")
    doc.add_heading("2 General description 总体概述", 1)
    doc.add_heading("2.1 Soft perspective 软件概述", 2)
    add_common_project_intro(doc)
    doc.add_heading("2.1.1 About the Project 项目介绍", 3)
    add_paragraph(doc, "系统通过后台管理界面集中处理企业办公事项。管理员维护基础数据并管理全局业务，员工完成个人办公事务，审批人员处理待办审批。系统通过站内信、WebSocket 和邮件提醒提高协同效率。")
    doc.add_heading("2.2 Soft function 软件功能", 2)
    add_image(doc, diagrams["系统功能模块图"], "图 2.1 软件功能结构图")
    doc.add_heading("2.3 User Roles and Permissions 用户角色与权限", 2)
    add_table(doc, ["角色", "可用功能", "限制"], [
        ("管理员", "用户、部门、审批、日程、新闻、字典、文件、通知、AI 日志等管理。", "仍需登录并携带有效 JWT。"),
        ("普通员工", "个人资料、发起审批、日程、新闻查看互动、站内信、智能问答。", "不能访问管理员接口，不能审批未分配给自己的审批单。"),
        ("审批员工", "在员工能力基础上处理待我审批和审批 AI 辅助。", "审批范围受 approver_id 和 is_approver 控制。"),
    ], [3, 9, 6])
    doc.add_heading("3 Functional Requirements 功能需求", 1)
    add_table(doc, ["模块", "需求描述"], [(m["title"], m["functions"]) for m in MODULES], [5, 11])
    for i, m in enumerate(MODULES, 1):
        doc.add_heading(f"3.{i} {m['title']}", 2)
        add_paragraph(doc, m["functions"])
        add_paragraph(doc, "主要接口：" + m["api"])
        add_paragraph(doc, "数据依据：" + m["tables"])
    doc.add_heading("4 Performance Requirements 非功能需求", 1)
    doc.add_heading("4.1 UI Requirements 界面要求", 2)
    add_bullets(doc, [
        "使用统一后台管理布局，包含侧边栏、顶部导航、面包屑、头像菜单和通知铃铛。",
        "表格支持分页、筛选、刷新和状态展示，危险操作需要二次确认。",
        "页面在桌面宽度下保持清晰，部分页面已有移动端截图用于适配检查。",
    ])
    add_image(doc, SCREENSHOTS["dashboard"], "图 4.1 首页仪表盘真实截图")
    doc.add_heading("4.2 Development Environment 开发环境", 2)
    add_table(doc, ["类别", "要求"], TECH_STACK, [5, 11])
    doc.add_heading("4.3 Development Rules 开发规范", 2)
    add_bullets(doc, [
        "后端接口统一 /api 前缀，返回 Result<T>，异常由 GlobalExceptionHandler 统一处理。",
        "业务表使用 created_at、updated_at、deleted 字段，删除采用逻辑删除。",
        "密码必须使用 BCrypt 加密，业务接口默认需要 JWT。",
        "高频数据使用 Redis 缓存，通知类事件通过 RabbitMQ 异步处理。",
        "前端 API 统一封装在 frontend/src/api，登录态由 Pinia 和本地 Token 管理。",
    ])
    doc.add_heading("5 Demand Classification 需求分级", 1)
    rows = []
    for idx, m in enumerate(MODULES, 1):
        priority = "A" if idx in [1, 2, 3, 4, 5, 9] else "B"
        rows.append((f"REQ-{idx:02d}", m["title"], priority))
    add_table(doc, ["Requirement ID 需求 ID", "Requirement Name 需求名称", "Classification 需求分级"], rows, [4, 9, 4])
    path = OUT / "企业OA管理系统软件需求说明书—程序员.docx"
    doc.save(path)
    return path


def build_design_doc(diagrams: dict[str, Path]) -> Path:
    doc = Document()
    style_doc(doc)
    add_cover(doc, "企业 OA 管理系统系统设计说明书", "System Design Specification")
    add_toc(doc, DESIGN_TOC)
    doc.add_page_break()
    doc.add_heading("1 Introduction  简介", 1)
    doc.add_heading("1.1 Purpose  文档目的", 2)
    add_paragraph(doc, "本文档从概要设计和模块详细设计两个层面说明 OA_SYSTEM 的系统结构、数据库、核心类、主要方法和模块实现边界。")
    doc.add_heading("1.2 Scope  本文档适用范围", 2)
    add_paragraph(doc, "适用于企业 OA 管理系统的设计说明、代码维护和课程验收。")
    doc.add_heading("1.3 Name 软件名称", 2)
    add_paragraph(doc, f"软件名称：{PROJECT_CN}，项目标识：{PROJECT_EN}。")
    doc.add_heading("1.4 Applications 软件应用领域", 2)
    add_paragraph(doc, "系统应用于企业内部办公管理场景，包括员工资料维护、组织管理、审批流转、会议日程、新闻公告、通知提醒和智能辅助办公。")
    doc.add_heading("2 High Level Design 概要设计", 1)
    doc.add_heading("2.1 系统功能设计", 2)
    add_image(doc, diagrams["系统功能模块图"], "图 2.1 系统功能设计图")
    doc.add_heading("2.2 系统架构设计", 2)
    add_image(doc, diagrams["系统总体架构图"], "图 2.2 系统架构设计图")
    for title, text in [
        ("2.2.1 前后端分离架构", "前端基于 Vue3 构建 SPA 管理端，后端基于 Spring Boot 提供 REST API，二者通过 Axios 和 JSON 数据交互。"),
        ("2.2.2 认证与权限架构", "系统使用 Spring Security + JWT 实现无状态认证，前端路由 meta 和后端 @PreAuthorize 共同完成权限控制。"),
        ("2.2.3 通知、RabbitMQ 与 WebSocket 链路", "业务事件进入 RabbitMQ，消费者生成站内信并推送 WebSocket，日程和系统通知可发送邮件。"),
    ]:
        doc.add_heading(title, 3)
        add_paragraph(doc, text)
    doc.add_heading("2.3 Database 数据库设计", 2)
    doc.add_heading("2.3.1 表关系", 3)
    add_image(doc, diagrams["数据库 ER 概要图"], "图 2.3 数据库表关系图")
    doc.add_heading("2.3.2 数据表设计", 3)
    add_table(doc, ["表名", "功能说明", "关键字段"], TABLES, [4, 4, 9])
    doc.add_heading("3 Low Level Model Design 模块详细设计", 1)
    for i, m in enumerate(MODULES, 1):
        doc.add_heading(f"3.{i} {m['title']}", 2)
        add_paragraph(doc, "(1) Overview 简介")
        add_paragraph(doc, m["functions"])
        add_paragraph(doc, "(2) Process 流程")
        add_paragraph(doc, m["flow"])
        add_paragraph(doc, "(3) Core Classes 核心类")
        add_table(doc, ["设计项", "内容"], [
            ("核心类/接口", m["classes"]),
            ("主要接口", m["api"]),
            ("数据表", m["tables"]),
            ("前端页面", m["frontend"]),
        ], [4, 12])
        add_paragraph(doc, "(4) Design Summary 设计小结")
        add_paragraph(doc, m["summary"])
    path = OUT / "企业OA管理系统系统设计说明书—程序员.docx"
    doc.save(path)
    return path


def build_all() -> list[Path]:
    ensure_dirs()
    diagrams = make_diagrams()
    docs = [
        build_course_report(diagrams),
        build_initiation_doc(diagrams),
        build_requirements_doc(diagrams),
        build_design_doc(diagrams),
    ]
    manifest = OUT / "文档生成清单.txt"
    lines = ["生成文档："] + [str(p) for p in docs] + ["", "生成图片："] + [str(p) for p in sorted(ASSETS.glob("*.png"))]
    manifest.write_text("\n".join(lines), encoding="utf-8")
    return docs


if __name__ == "__main__":
    for doc_path in build_all():
        print(doc_path)
